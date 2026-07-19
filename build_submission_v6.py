#!/usr/bin/env python3
"""Build the MA-ABE-FU v6 IEEE submission package."""

from __future__ import annotations

import csv
import json
import re
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "submission_tifs_v6"
FIG = OUT / "figure"
REPRO = OUT / "reproducibility"
PUBLIC_DATA = ROOT / "public_data"
REPRO_REPO = ROOT / "ma-abe-fu-reproducibility"
FINAL_LAYOUT = ROOT / "final_layout"
SUBMISSION_REQUIRED = FINAL_LAYOUT / "submission_required"
SUPPORTING_FILES = FINAL_LAYOUT / "supporting_files"

TITLE = "MA-ABE-FU: Policy-Authenticated and Evidence-Bound Federated Unlearning for Cross-Border Identity Authentication"
AUTHORS = "Jian Chen, Sheng Peng, Zhiming Cai, Jiayin Qi, and Fu Mo"
EMAIL = "mofu@gdust.edu.cn"
TEL = "+86 15989698699"
TARGET = "IEEE Transactions on Information Forensics and Security"
CODE_URL = "https://github.com/Kent919/MA-ABE-FU"

REFS = [
    "Financial Action Task Force, \"Guidance on Digital Identity,\" Paris, France, 2020. [Online]. Available: https://www.fatf-gafi.org/en/publications/Financialinclusionandnpoissues/Digital-identity-guidance.html",
    "European Banking Authority, \"Guidelines on the use of remote customer onboarding solutions,\" EBA/GL/2022/15, Nov. 22, 2022. [Online]. Available: https://www.eba.europa.eu/regulation-and-policy/anti-money-laundering-and-countering-financing-terrorism/guidelines-use-remote-customer-onboarding-solutions",
    "European Parliament and Council of the European Union, Regulation (EU) 2016/679, General Data Protection Regulation, OJ L 119, pp. 1-88, May 4, 2016, arts. 5(2), 17, 30, and 32. [Online]. Available: https://eur-lex.europa.eu/eli/reg/2016/679/oj",
    "European Commission, Commission Implementing Decision (EU) 2021/914 on standard contractual clauses for the transfer of personal data to third countries, OJ L 199, pp. 31-61, Jun. 7, 2021, clauses 8 and 15. [Online]. Available: https://eur-lex.europa.eu/eli/dec_impl/2021/914/oj",
    "Standing Committee of the National People's Congress, Personal Information Protection Law of the People's Republic of China, Aug. 20, 2021, arts. 38-40, 44, and 47. [Online]. Available: https://www.npc.gov.cn/",
    "California State Legislature, California Consumer Privacy Act of 2018, Cal. Civ. Code Sec. 1798.105, 2018, as amended. [Online]. Available: https://leginfo.legislature.ca.gov/",
    "B. McMahan, E. Moore, D. Ramage, S. Hampson, and B. A. y Arcas, \"Communication-Efficient Learning of Deep Networks From Decentralized Data,\" in Proc. AISTATS, PMLR, vol. 54, pp. 1273-1282, 2017.",
    "P. Kairouz et al., \"Advances and Open Problems in Federated Learning,\" Found. Trends Mach. Learn., vol. 14, no. 1-2, pp. 1-210, 2021, doi: 10.1561/2200000083.",
    "G. Liu, X. Ma, Y. Yang, C. Wang, and J. Liu, \"FedEraser: Enabling Efficient Client-Level Data Removal From Federated Learning Models,\" in Proc. IEEE/ACM IWQoS, pp. 1-10, 2021, doi: 10.1109/IWQOS52092.2021.9521274.",
    "L. Zhang, T. Zhu, H. Zhang, P. Xiong, and W. Zhou, \"FedRecovery: Differentially Private Machine Unlearning for Federated Learning Frameworks,\" IEEE Trans. Inf. Forensics Secur., vol. 18, pp. 4732-4746, 2023, doi: 10.1109/TIFS.2023.3297905.",
    "X. Gao, X. Ma, J. Wang, Y. Sun, B. Li, S. Ji, P. Cheng, and J. Chen, \"VeriFi: Towards Verifiable Federated Unlearning,\" IEEE Trans. Dependable Secure Comput., vol. 21, no. 6, pp. 5720-5736, Nov. 2024, doi: 10.1109/TDSC.2024.3382321.",
    "Z. Liu, Y. Jiang, W. Jiang, J. Guo, J. Zhao, and K.-Y. Lam, \"Guaranteeing Data Privacy in Federated Unlearning With Dynamic User Participation,\" IEEE Trans. Dependable Secure Comput., vol. 22, no. 3, pp. 2072-2085, May 2025, doi: 10.1109/TDSC.2024.3476533.",
    "J. Chen, Z. Lin, W. Lin, W. Shi, X. Yin, and D. Wang, \"FedMUA: Exploring the Vulnerabilities of Federated Learning to Malicious Unlearning Attacks,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 1665-1678, 2025, doi: 10.1109/TIFS.2025.3531141.",
    "Z. Liu, H. Ye, Y. Jiang, J. Shen, J. Guo, I. Tjuawinata, and K.-Y. Lam, \"Privacy-Preserving Federated Unlearning With Certified Client Removal,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 3966-3978, 2025, doi: 10.1109/TIFS.2025.3555868.",
    "S. Zhao, J. Zhang, X. Ma, Q. Jiang, Z. Ma, S. Gao, Z. Ying, and J. Ma, \"FedWiper: Federated Unlearning via Universal Adapter,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 4042-4054, 2025, doi: 10.1109/TIFS.2025.3557671.",
    "K. Gao, T. Zhu, D. Ye, B. Liu, and W. Zhou, \"Hidden Threats in Federated Unlearning: Camouflaged Poisoning Attacks and Their Unlearning Consequences,\" IEEE Trans. Dependable Secure Comput., vol. 23, no. 2, pp. 2934-2948, Mar. 2026, doi: 10.1109/TDSC.2025.3630811.",
    "L. Bourtoule et al., \"Machine Unlearning,\" in Proc. IEEE Symp. Security and Privacy, pp. 141-159, 2021, doi: 10.1109/SP40001.2021.00019.",
    "C. Guo, T. Goldstein, A. Hannun, and L. van der Maaten, \"Certified Data Removal From Machine Learning Models,\" in Proc. ICML, PMLR, vol. 119, pp. 3832-3842, 2020.",
    "S. Garg, S. Goldwasser, and P. N. Vasudevan, \"Formalizing Data Deletion in the Context of the Right to Be Forgotten,\" in EUROCRYPT, LNCS, vol. 12106, pp. 373-402, 2020, doi: 10.1007/978-3-030-45724-2_13.",
    "Z. Liu, Y. Jiang, J. Shen, M. Peng, K.-Y. Lam, X. Yuan, and X. Liu, \"A Survey on Federated Unlearning: Challenges, Methods, and Future Directions,\" ACM Comput. Surv., vol. 57, no. 1, pp. 1-38, Jan. 2025, doi: 10.1145/3679014.",
    "T. T. Nguyen et al., \"A Survey of Machine Unlearning,\" ACM Trans. Intell. Syst. Technol., vol. 16, no. 5, Art. 108, 2025, doi: 10.1145/3749987.",
    "F. Wang, B. Li, and B. Li, \"Federated Unlearning and Its Privacy Threats,\" IEEE Network, vol. 38, no. 2, pp. 294-300, Mar.-Apr. 2024, doi: 10.1109/MNET.004.2300056.",
    "A. Sahai and B. Waters, \"Fuzzy Identity-Based Encryption,\" in EUROCRYPT, LNCS, vol. 3494, pp. 457-473, 2005, doi: 10.1007/11426639_27.",
    "V. Goyal, O. Pandey, A. Sahai, and B. Waters, \"Attribute-Based Encryption for Fine-Grained Access Control of Encrypted Data,\" in Proc. ACM CCS, pp. 89-98, 2006, doi: 10.1145/1180405.1180418.",
    "J. Bethencourt, A. Sahai, and B. Waters, \"Ciphertext-Policy Attribute-Based Encryption,\" in Proc. IEEE Symp. Security and Privacy, pp. 321-334, 2007, doi: 10.1109/SP.2007.11.",
    "M. Chase, \"Multi-Authority Attribute Based Encryption,\" in TCC, LNCS, vol. 4392, pp. 515-534, 2007, doi: 10.1007/978-3-540-70936-7_28.",
    "A. Lewko and B. Waters, \"Decentralizing Attribute-Based Encryption,\" in EUROCRYPT, LNCS, vol. 6632, pp. 568-588, 2011, doi: 10.1007/978-3-642-20465-4_31.",
    "B. Waters, \"Ciphertext-Policy Attribute-Based Encryption: An Expressive, Efficient, and Provably Secure Realization,\" in PKC, LNCS, vol. 6571, pp. 53-70, 2011, doi: 10.1007/978-3-642-19379-8_4.",
    "Y. Zhang, R. H. Deng, S. Xu, J. Sun, Q. Li, and D. Wu, \"Attribute-Based Encryption for Cloud Computing Access Control: A Survey,\" ACM Comput. Surv., vol. 53, no. 4, pp. 1-41, 2020, doi: 10.1145/3398036.",
    "X. Xing, Y. Liu, Q. Wu, Z. Guan, D. Li, D. Li, Y. Lu, and W. Susilo, \"Multi-Committee ABE Based Decentralized Access Control With Sharding Blockchain for Web 3.0,\" IEEE Trans. Dependable Secure Comput., vol. 22, no. 3, pp. 2533-2549, May 2025, doi: 10.1109/TDSC.2024.3520121.",
    "D. Boneh and M. Franklin, \"Identity-Based Encryption From the Weil Pairing,\" SIAM J. Comput., vol. 32, no. 3, pp. 586-615, 2003, doi: 10.1137/S0097539701398521.",
    "A. Fiat and A. Shamir, \"How to Prove Yourself: Practical Solutions to Identification and Signature Problems,\" in CRYPTO, LNCS, vol. 263, pp. 186-194, 1987, doi: 10.1007/3-540-47721-7_12.",
    "S. Goldwasser, S. Micali, and C. Rackoff, \"The Knowledge Complexity of Interactive Proof Systems,\" SIAM J. Comput., vol. 18, no. 1, pp. 186-208, 1989, doi: 10.1137/0218012.",
    "J. Groth, \"On the Size of Pairing-Based Non-Interactive Arguments,\" in EUROCRYPT, LNCS, vol. 9666, pp. 305-326, 2016, doi: 10.1007/978-3-662-49896-5_11.",
    "B. Bunz et al., \"Bulletproofs: Short Proofs for Confidential Transactions and More,\" in Proc. IEEE Symp. Security and Privacy, pp. 315-334, 2018, doi: 10.1109/SP.2018.00020.",
    "A. Gabizon, Z. J. Williamson, and O. Ciobotaru, \"PLONK: Permutations Over Lagrange-Bases for Oecumenical Noninteractive Arguments of Knowledge,\" Cryptology ePrint Archive, Rep. 2019/953, 2019. [Online]. Available: https://eprint.iacr.org/2019/953",
    "E. Ben-Sasson, I. Bentov, Y. Horesh, and M. Riabzev, \"Scalable, Transparent, and Post-Quantum Secure Computational Integrity,\" Cryptology ePrint Archive, Rep. 2018/046, 2018. [Online]. Available: https://eprint.iacr.org/2018/046",
    "E. Androulaki et al., \"Hyperledger Fabric: A Distributed Operating System for Permissioned Blockchains,\" in Proc. EuroSys, Art. no. 30, pp. 30:1-30:15, 2018, doi: 10.1145/3190508.3190538.",
    "S. Nakamoto, \"Bitcoin: A Peer-to-Peer Electronic Cash System,\" 2008. [Online]. Available: https://bitcoin.org/bitcoin.pdf",
    "H. Hofmann, \"Statlog (German Credit Data),\" UCI Machine Learning Repository, 1994, doi: 10.24432/C5NC77.",
    "S. Moro, P. Rita, and P. Cortez, \"Bank Marketing,\" UCI Machine Learning Repository, 2014, doi: 10.24432/C5K306.",
    "S. Jesus, J. Pombal, D. Alves, A. Cruz, P. Saleiro, R. Ribeiro, J. Gama, and P. Bizarro, \"Turning the Tables: Biased, Imbalanced, Dynamic Tabular Datasets for ML Evaluation,\" in Adv. Neural Inf. Process. Syst. 35, pp. 33563-33575, 2022, doi: 10.52202/068431-2432.",
    "R. Shokri, M. Stronati, C. Song, and V. Shmatikov, \"Membership Inference Attacks Against Machine Learning Models,\" in Proc. IEEE Symp. Security and Privacy, pp. 3-18, 2017, doi: 10.1109/SP.2017.41.",
    "S. Yeom, I. Giacomelli, M. Fredrikson, and S. Jha, \"Privacy Risk in Machine Learning: Analyzing the Connection to Overfitting,\" in Proc. IEEE CSF, pp. 268-282, 2018, doi: 10.1109/CSF.2018.00027.",
]

BIOS = [
    "Jian Chen, PhD. is a senior researcher at Macau Millennium College and a member of the Macau-Qin'ao Cross-Border Data Flow Research Team. His research interests include smart cities and digital security.",
    "Jiayin Qi, PhD. is a professor of Cyberspace Security at Guangzhou University and executive director of the Institute for Innovation in Cyberspace Governance, Guangdong-Hong Kong-Macao Greater Bay Area. Her research interests include cross-border data governance, privacy-preserving computation, and trustworthy data sharing.",
    "Sheng Peng, PhD. is an associate professor at Macau Millennium College and a postdoctoral research fellow. His research interests include cross-border digital identity, large-scale system implementation, and data security engineering.",
    "Zhiming Cai, PhD. is a professor at Macau Millennium College. His research interests include cross-border digital identity, large-scale system implementation, and data security engineering.",
    "Fu Mo. PhD. is an associate professor at The Institute of Data Intelligence, Guangdong University of Science and Technology, where he also serves as deputy director of the research affairs office. His research interests include intelligent electronic systems and data analytics, with particular focus on big-data analysis and decision-support technologies.",
]


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.70)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.68)
    sec.right_margin = Inches(0.68)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(10.0)
    doc.styles["Heading 1"].font.size = Pt(12.0)
    doc.styles["Heading 2"].font.size = Pt(11.0)
    doc.styles["Heading 3"].font.size = Pt(10.5)


def set_columns(section, count: int, space_twips: int = 252) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def set_doc_flag(doc: Document, name: str, value: bool) -> None:
    setattr(doc, name, value)


def two_column_enabled(doc: Document) -> bool:
    return bool(getattr(doc, "_maabe_two_columns", False))


def in_wide_block(doc: Document) -> bool:
    return bool(getattr(doc, "_maabe_wide_block", False))


def start_two_column_body(doc: Document) -> None:
    set_columns(doc.sections[-1], 1)
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec, 2)
    set_doc_flag(doc, "_maabe_two_columns", True)
    set_doc_flag(doc, "_maabe_wide_block", False)


def start_wide_block(doc: Document) -> None:
    if not two_column_enabled(doc) or in_wide_block(doc):
        return
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec, 1)
    set_doc_flag(doc, "_maabe_wide_block", True)


def end_wide_block(doc: Document) -> None:
    if not two_column_enabled(doc) or not in_wide_block(doc):
        return
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec, 2)
    set_doc_flag(doc, "_maabe_wide_block", False)


def para(doc: Document, text: str = "", bold: bool = False, italic: bool = False, center: bool = False, size: float = 10.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.0)
    p.paragraph_format.line_spacing = 1.04
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    return p


def formula(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for idx, line in enumerate(str(text).split("\n")):
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        r.font.size = Pt(8.4 if two_column_enabled(doc) else 9.3)
        if idx < len(str(text).split("\n")) - 1:
            r.add_break()


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def looks_numeric(value) -> bool:
    try:
        float(str(value).strip())
        return True
    except ValueError:
        return False


def cell_text(cell, value, bold: bool = False, size: float = 8.0, align=None) -> None:
    cell.text = ""
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    if align is None:
        align = WD_ALIGN_PARAGRAPH.CENTER if looks_numeric(value) else WD_ALIGN_PARAGRAPH.LEFT
    p.alignment = align
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_caption(doc: Document, num: str, title: str):
    start_wide_block(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"TABLE {num}\n{title.upper()}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)


def add_table(doc: Document, headers, rows, widths=None, size: float = 7.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(t.rows[0].cells[i], "EAF0F7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, False, size=size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    para(doc)
    end_wide_block(doc)
    return t


def add_fig(doc: Document, filename: str, caption: str, width: float = 6.92):
    start_wide_block(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(7)
    r = c.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)
    end_wide_block(doc)


def alg_box(doc: Document, title: str, body: str):
    start_wide_block(doc)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, "F6F8FB")
    set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rr = p2.add_run(body)
    rr.font.name = "Courier New"
    rr._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    rr.font.size = Pt(7.8)
    para(doc)
    end_wide_block(doc)


def appendix_proof_map(doc: Document) -> None:
    start_wide_block(doc)
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    heads = ["G0", "G1", "G2", "G3", "G4", "G5"]
    bodies = [
        "real experiment",
        "random-oracle commitments",
        "MA-ABE capsule swap",
        "padded channel view",
        "UCAP signature and chain",
        "ZK soundness",
    ]
    for i, head in enumerate(heads):
        cell_text(table.rows[0].cells[i], head, True, size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i], "EAF0F7")
        cell_text(table.rows[1].cells[i], bodies[i], False, size=7.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.04)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(7)
    r = c.add_run("Fig. A1. Reduction roadmap. Each hop hides or authenticates one component of the adversary's view.")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.6)
    end_wide_block(doc)


def load_csv(name: str):
    with open(REPRO / name, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fnum(value, digits=3) -> str:
    return f"{float(value):.{digits}f}"


def mean(rows, filters, col: str) -> float:
    vals = [float(r[col]) for r in rows if all(r[k] == v for k, v in filters.items())]
    return sum(vals) / len(vals)


def primary_rows(raw):
    methods = ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU"]
    rows = []
    for dataset in ["German Credit", "Bank Marketing"]:
        for method in methods:
            filt = {"dataset": dataset, "method": method}
            rows.append(
                [
                    dataset,
                    method,
                    fnum(mean(raw, filt, "retained_auc"), 3),
                    fnum(mean(raw, filt, "mia_gap"), 3),
                    fnum(mean(raw, filt, "l2_to_oracle"), 3),
                    fnum(mean(raw, filt, "runtime_to_oracle"), 3),
                    fnum(mean(raw, filt, "attribute_leak_auc"), 3),
                ]
            )
    return rows


def leakage_rows(raw):
    methods = ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU"]
    rows = []
    for method in methods:
        filt = {"method": method}
        rows.append(
            [
                method,
                fnum(mean(raw, filt, "type_leak_auc"), 3),
                fnum(mean(raw, filt, "attribute_leak_auc"), 3),
                fnum(mean(raw, filt, "mia_gap"), 3),
            ]
        )
    return rows


def proxy_rows(ablation):
    rows = []
    for r in ablation:
        if r["ablation"] != "full proxy":
            continue
        rows.append(
            [
                r["dataset"],
                r["method"].replace("-proxy", ""),
                fnum(r["mean_abs_auc_delta_to_oracle"], 4),
                fnum(r["mean_mia_gap_delta_to_oracle"], 4),
                fnum(r["mean_runtime_to_oracle"], 3),
                fnum(r["max_abs_auc_delta_to_oracle"], 4),
            ]
        )
    return rows


def crypto_rows(crypto):
    rows = []
    for r in crypto:
        backend = r["backend"]
        policy_rows = int(float(r["policy_rows"]))
        auth = int(float(r["authority_count"]))
        keep = backend == "primitive_modexp_proxy" and policy_rows in {4, 48}
        keep = keep or (backend == "bn254_pairing_py_ecc" and policy_rows in {4, 48} and auth in {2, 4, 6})
        if not keep:
            continue
        capsule = r["ma_abe_capsule_proxy_ms"] if backend == "primitive_modexp_proxy" else r["bn254_pairing_ma_abe_ms"]
        rows.append(
            [
                "primitive proxy" if backend == "primitive_modexp_proxy" else "BN254 pairing",
                "-" if auth == 0 else str(auth),
                str(policy_rows),
                fnum(capsule, 1),
                fnum(r["padded_envelope_hmac_ms"], 3),
                fnum(r["ucap_rsa_pss_sign_ms"], 3),
                fnum(r["total_control_plane_ms"], 1),
            ]
        )
    return rows


def risk_rows(riskgrid):
    rows = []
    for r in riskgrid:
        if float(r["beta"]) == 2.0 and float(r["alpha"]) in {0.0, 0.25, 0.5} and r["method"] in {"FedEraser-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"}:
            rows.append(
                [
                    r["method"].replace("-proxy", ""),
                    fnum(r["alpha"], 2),
                    fnum(r["mean_riskgap"], 3),
                    fnum(r["audit_pass_rate"], 2),
                    fnum(r["mean_model_residue"], 3),
                ]
            )
    return rows


def add_introduction(doc, raw):
    doc.add_heading("I. Introduction", level=1)
    para(doc, "Cross-border identity authentication is now a distributed-systems problem. A bank may onboard a mobile customer in one jurisdiction while relying on credit-bureau attributes, telecom evidence, sanctions screening, and fraud signals from others. The same institution may train a federated risk model across regional branches, each subject to local transfer and customer-due-diligence rules. Digital identity guidance requires risk-based identity proofing [1]. Remote-onboarding guidance requires accountable controls for non-face-to-face verification [2]. Data-protection law gives subjects rights that remain relevant even when model training is federated [3]-[6].")
    para(doc, "A forget request in this setting is not only a model update. The server must know that the request was authorized by the right authorities. It should not learn sensitive policy attributes from the request surface. An auditor should later verify the event without receiving the full policy. In credit and banking identity checks, even the reason for a forget request can be sensitive, such as a foreign-worker flag, a personal-loan attribute, or a high-risk onboarding channel.")
    para(doc, "Existing work leaves three gaps. First, federated learning and early federated unlearning methods focus on the optimizer. They update a model after a forget request, but they do not bind the repair to a multi-authority authorization predicate [7]-[10]. Second, recent verifiable and privacy-preserving federated unlearning improves accountability, client churn handling, attack analysis, certified repair, and adapter-based repair. The visible request surface, however, is still largely controlled by the server and can leak request type or policy scope [11]-[16]. Third, machine unlearning theory, certified data repair, and SISA-style retraining clarify retained-set targets, but they usually assume a single requester or controller [17]-[22]. Cross-border identity workflows require several authorities to approve the same request without exposing the policy to one administrative domain.")
    para(doc, "MA-ABE-FU addresses this control gap. It does not use MA-ABE to repair a model. Instead, MA-ABE authorizes and hides the policy capsule that triggers a learning-side repair module. The repair module can be retraining, cached replay, recovery, certified repair, or another optimizer. The contribution is the envelope around that module: a padded request channel, proof relations for authorization and repair consistency, and a UCAP evidence object that binds policy, channel, model, and residual-risk commitments.")
    add_fig(doc, "Fig. 1.tif", "Fig. 1. MA-ABE-FU control plane. Multi-authority policy authorization, learning-side unlearning repair, and UCAP audit evidence are separated but transcript-bound.")
    para(doc, "This paper makes three contributions.")
    contribs = [
        "A cross-border unlearning protocol that turns a forget request into a policy-authenticated event before any repair is accepted. In the experiments, this control plane lowers mean attribute-leakage AUC from 0.671-0.907 for representative FU baselines to 0.503.",
        "A security formulation for policy-authenticated update hiding. The proof separates cryptographic indistinguishability from empirical model residue, leaving the residual term as an explicit audit parameter.",
        "A reproducible experimental package for non-IID public financial datasets. Across 126 federated runs and 16 proxy-ablation cells, MA-ABE-FU keeps average retained AUC at 0.787 and reports measured control-plane cost under primitive and BN254 backends.",
    ]
    for item in contribs:
        doc.add_paragraph(item, style="List Number")


def add_related_work(doc):
    doc.add_heading("II. Related Work", level=1)
    para(doc, "Federated unlearning mechanisms usually operate inside the learning plane. FedEraser, FedRecovery, VeriFi, dynamic-participation methods, FedMUA, certified repair, and FedWiper differ in how they replay, recover, certify, attack, or adapt a model [9]-[16]. They share a basic assumption: the server can observe enough of the request process to run the selected repair. This assumption is costly in cross-border identity authentication, where the request surface can reveal why the policy was invoked. MA-ABE-FU keeps these repair methods as possible modules but hides the policy capsule that selects them.")
    para(doc, "Machine unlearning theory gives the retained-set target and residual-risk language used to judge a repair [17]-[22]. Its authorization model is usually simple: a requester or controller asks for a record to be forgotten. Cross-border credit and banking workflows need a stricter rule. Several authorities may have to agree that the scope is valid, and their joint decision should not be exposed to the training server.")
    para(doc, "Attribute-based encryption provides the policy machinery for that decision. Prior work covers fine-grained encryption, CP-ABE, multi-authority issuance, decentralized ABE, and LSSS policy encoding [23]-[29]. Recent multi-committee ABE work shows that decentralized access control can scale under distributed administration [30]. Access control alone is still insufficient. It proves who may decrypt; it does not prove that an unlearning repair was run or that the request metadata remained policy-private.")
    para(doc, "Proof systems and append-only ledgers provide the audit substrate [31]-[39], but their value depends on the statement being proved. MA-ABE-FU binds authorization, padded channel shape, repair transcript, and residual risk in one UCAP object.")
    table_caption(doc, "I", "Mechanism gap in existing approaches")
    add_table(
        doc,
        ["Line of work", "Bottom mechanism", "Why it fails in cross-border multi-authorization", "MA-ABE-FU response"],
        [
            ["FU replay/recovery", "cached updates, recovery, adapters", "server-visible request type and policy scope remain outside the optimizer", "hide request capsule and bind repair to UCAP evidence"],
            ["Verifiable FU", "proof of an unlearning computation", "proof target does not encode multi-authority authorization privacy", "add pi_auth and padded update-family channel"],
            ["Certified repair/SISA", "shards or repair certificates", "single-domain request semantics do not express independent authorities", "compile forget scope as LSSS multi-authority policy"],
            ["MA-ABE/access control", "policy-based decryption", "access authorization does not measure model residue", "decouple authorization from learning-side repair"],
            ["Ledger/ZK audit", "append-only evidence and proof systems", "statements can omit malicious-server observability", "bind channel, policy, model commitments, and RiskGap"],
        ],
        [1.15, 1.45, 2.25, 1.65],
        size=7.4,
    )


def add_threat_model(doc):
    doc.add_heading("III. System and Threat Model", level=1)
    para(doc, "Entities. The system contains clients C_1,...,C_n, an authentication server AS, attribute authorities AA_1,...,AA_m, a requester Rq, and an auditor Aud. Each AA_i controls a disjoint or partially overlapping attribute domain U_i and issues GID-bound keys. A forget policy P is represented as an LSSS matrix with row map rho.")
    formula(doc, "P = (M, rho),    M in Z_p^{ell x k},    rho: {1,...,ell} -> union_i U_i.                         (1)")
    para(doc, "Planes. The control plane authorizes, hides, records, and audits the forget event. The learning plane receives an approved repair scope and runs a selected FU optimizer. No step in the protocol assumes that MA-ABE decrypts model parameters or performs statistical repair.")
    formula(doc, "F = (P, scope, tau, cfg),    S_F = {z in D : phi_P(z)=1 and epoch(z) <= tau}.                     (2)")
    para(doc, "Server behavior is split into two cases. During ordinary training the server is honest-but-curious: it executes FedAvg-style aggregation, observes timestamps, envelope lengths, client identifiers allowed by the deployment, update norms, and public configuration, and tries to infer membership or sensitive attributes. During the forget request phase it may be malicious: it can classify padded envelopes, infer policy attributes from metadata, delay unlearning repair, weaken repair rounds, omit clients from repair, reorder or suppress audit records, and forge evidence objects.")
    para(doc, "Attacker boundary. The adversary may adaptively request transcripts, corrupt a bounded set of clients, choose equal-length challenge envelopes, and query audit verification. It may not corrupt all authorities needed to satisfy the challenge policy, break MA-ABE collusion resistance, forge AS signatures, find commitment collisions, or make a false statement accepted by the proof system. Weak or incomplete repair is not hidden by the theorem; it is detected by pi_rep and the empirical RiskGap report.")
    table_caption(doc, "II", "Notation and formal objects")
    add_table(
        doc,
        ["Symbol", "Meaning", "Security role"],
        [
            ["AA_i, U_i, GID", "authority i, its attribute domain, and global identity", "multi-authority key boundary"],
            ["P=(M,rho)", "LSSS policy with matrix M in Z_p^{ell x k} and row map rho", "hidden authorization predicate"],
            ["F=(P,scope,tau,cfg)", "forget request descriptor: policy, scope, epoch, repair configuration", "request semantics"],
            ["S_F", "records satisfying phi_P within scope and epoch tau", "forget scope for unlearning repair"],
            ["Pad, e_F, ch", "padding mechanism, padded envelope, and channel transcript", "request-type hiding surface"],
            ["CT_F, tag", "MA-ABE capsule and encrypted request tag", "policy-bearing ciphertext"],
            ["c_P,c_s,c_cfg,h_F", "commitments to policy, scope, configuration, and request handle", "audit binding"],
            ["theta_t, theta'_t", "model commitments before and after unlearning repair", "repair binding"],
            ["R, s_l2, tau_R", "residual report, L2 normalization scale, and audit threshold", "RiskGap calibration"],
            ["R_auth, w_auth", "authorization proof relation and witness", "policy satisfaction without policy disclosure"],
            ["R_rep, w_rep, T_rep", "repair-consistency relation, witness, and repair trace", "unlearning repair consistency"],
            ["pi_auth, pi_rep", "zero-knowledge proofs for R_auth and R_rep", "proof handles"],
            ["O_F", "UCAP evidence object with commitments, proofs, signature, and chain pointer", "audit-chain unit"],
            ["alpha,beta", "RiskGap weights for model residue and utility loss", "residual-risk calibration"],
        ],
        [0.95, 2.45, 2.75],
        size=6.7,
    )


def add_protocol(doc):
    doc.add_heading("IV. MA-ABE-FU Protocol", level=1)
    para(doc, "The construction exposes four algorithms. The first two establish a policy-authenticated forget request. The third performs evidence-bound repair through a learning-side module. The fourth verifies the evidence object. All ciphertexts, commitments, and proofs are bound to a deployment identifier, an epoch, and a repair configuration to prevent cross-run replay.")
    alg_box(doc, "Algorithm 1: SystemSetup(1^lambda, {AA_i}, U, cfg)", "Input : security parameter, authorities, attribute domains, configuration.\nOutput: public parameters pp and local authority secrets.\n1: for each authority AA_i do\n2:     sample (mpk_i, msk_i) over U_i.\n3:     issue GID-bound keys SK_{GID,i,S_i}.\n4: AS samples signing key sk_AS and sets h_0 = H(domain || cfg).\n5: publish pp = ({mpk_i}, Com, Pad, ZK parameters, grammar, tau_R).")
    alg_box(doc, "Algorithm 2: ForgetRequest(pp, GID, P, scope, tau, cfg)", "Input : policy P, request scope, cutoff epoch tau, repair configuration.\nOutput: padded forget envelope e_F.\n1: compile P to LSSS matrix (M,rho).\n2: tag <- H(GID || scope || tau || cfg || nonce).\n3: CT_F <- MA-ABE.Enc({mpk_i}, tag; P).\n4: commit c_P=H(P), c_s=H(scope), c_cfg=H(cfg), h_F=H(CT_F || c_s || tau).\n5: prove pi_auth for R_auth without revealing satisfying attributes.\n6: emit e_F <- Pad(type=forget, CT_F, c_P, c_s, c_cfg, tau, pi_auth).")
    alg_box(doc, "Algorithm 3: EvidenceBoundRepair(pp, e_F, theta_t, clients)", "Input : padded envelope, current model, participating clients.\nOutput: repaired model theta'_t, residual report R, evidence O_F.\n1: check padded grammar and CT_F well-formedness.\n2: verify pi_auth; reject before repair if it fails.\n3: derive S_F from the authorized scope.\n4: theta'_t, T_rep <- Repair(theta_t, D \\ S_F, cfg).\n5: commit ch, theta_t, theta'_t, T_rep, and residual report R.\n6: prove pi_rep for R_rep and append O_F to the UCAP chain.")
    alg_box(doc, "Algorithm 4: AuditVerify(pp, O_F, tau_R)", "Input : evidence object and audit threshold tau_R.\nOutput: accept or reject.\n1: verify sigma_AS and predecessor hash h_prev.\n2: verify pi_auth on (c_P,c_s,c_cfg,CT_F,h_F,epoch).\n3: verify pi_rep on (H(ch),H(theta_t),H(theta'_t),H(R),cfg).\n4: recompute or sample-check R on held-out audit data.\n5: accept iff all checks pass and RiskGap(R) <= tau_R.")
    para(doc, "The two proof relations are the cryptographic boundary between hidden authorization and observable learning repair. Let w_auth contain a satisfying attribute-key tuple and encryption randomness. Let w_rep contain the repair trace, retained-client commitments, and opening randomness for model and report commitments.")
    formula(doc, "R_auth: MA-ABE.Open(CT_F,w_auth)=tag and phi_P(S_GID)=1.                                      (3)")
    formula(doc, "R_rep: T_rep follows Repair(cfg,D\\S_F) and opens commitments to theta'_t and R.                  (4)")
    formula(doc, "pi_auth <- ZK.Prove(R_auth; w_auth),    pi_rep <- ZK.Prove(R_rep; w_rep).                         (5)")
    para(doc, "The server therefore cannot use plaintext policy parsing as the repair trigger. It observes a padded envelope family and commitments; policy satisfaction is checked through pi_auth. Conversely, pi_rep does not prove that all influence has vanished. It proves that the declared repair transcript is consistent with the committed configuration, after which empirical residual risk is checked separately.")
    formula(doc, "RiskGap(R) = |AUC_MIA - 0.5| + alpha * ||theta'_t - theta_oracle||_2 / s_l2 + beta * max(0,U_oracle - U_ret).        (6)")


def add_security(doc):
    doc.add_heading("V. Security Analysis", level=1)
    para(doc, "Definition 1, policy-authenticated update hiding. The adversary A receives pp and may query key-issue, transcript, repair, and audit oracles, except that it may not obtain enough challenge attributes to satisfy P*. A chooses two equal-public-shape events e_0 and e_1: e_0 is an ordinary training update and e_1 is a valid forget event under P*. The challenger samples b, returns View_b, and A wins if it outputs b'=b.")
    formula(doc, "Adv_auth(A) = | Pr[b'=b] - 1/2 |.                                                              (7)")
    para(doc, "The game covers the malicious-server actions stated in the threat model. Metadata classification is the View_b distinction problem. Attribute inference is blocked by the MA-ABE capsule and pi_auth. Audit-chain forgery is reduced to signature and commitment failure. Weak repair is not modeled away; it must either violate pi_rep or appear as residual RiskGap.")
    para(doc, "Theorem 1. Assume selective collusion-resistant MA-ABE, programmable random-oracle commitments, computationally indistinguishable padding for equal-shape envelopes, EUF-CMA secure AS signatures, and knowledge soundness for pi_auth and pi_rep. For every PPT adversary A against Definition 1 there exist PPT adversaries B_1,...,B_5 such that")
    formula(doc, "Adv_auth(A) <= Adv_MA-ABE(B_1) + Adv_RO(B_2) + Adv_PAD(B_3) + Adv_SIG(B_4) + Adv_ZK(B_5) + eps_R + negl(lambda).       (8)")
    para(doc, "The term eps_R is the configured residual-risk tolerance. It is explicit because model influence is empirical for the selected learning module. This separation is important: the theorem proves policy-authenticated update hiding and audit binding for the control plane, not unconditional disappearance of model influence by encryption.")
    add_fig(doc, "Fig. 2.tif", "Fig. 2. Policy-authenticated update-hiding game and reduction purpose. G0-G5 respectively bound commitment replacement, MA-ABE capsule indistinguishability, padded-channel request-type hiding, UCAP-chain unforgeability, and zero-knowledge proof soundness.")
    para(doc, "Proof sketch. The proof uses a standard game sequence. Each hop replaces one observable component in the server view or turns a successful forgery into a primitive-level adversary. The final game contains only padded metadata, random commitments, valid audit evidence, and the residual-risk report. The remaining advantage is therefore the configured tolerance eps_R plus negligible terms. Appendix A gives the step-by-step reduction.")


def add_experiments(doc, meta, bafs_status):
    doc.add_heading("VI. Experimental Methodology", level=1)
    para(doc, "Data and partitioning. Experiments use two public financial identity datasets: UCI German Credit and UCI Bank Marketing [40], [41]. German Credit is partitioned into 12 clients and Bank Marketing into 24 clients. Client partitions are Dirichlet non-IID with alpha=0.35, and each dataset is evaluated over three random seeds and three forget ratios: 0.25, 0.50, and 1.00 of the policy-eligible records. The BAFS suite is implemented as an optional loader because it is the preferred public fraud benchmark for large-scale banking identity experiments [42]. In this run, no authenticated BAFS CSV was present locally; the package records this status and does not report BAFS metrics.")
    table_caption(doc, "III", "Datasets and federated partition settings")
    add_table(
        doc,
        ["Dataset", "Rows", "Encoded features", "Clients", "Rounds", "Forget policy"],
        [
            ["German Credit", meta["German Credit"]["instances"], meta["German Credit"]["encoded_features"], meta["German Credit"]["clients"], meta["German Credit"]["rounds"], meta["German Credit"]["forget_policy"]],
            ["Bank Marketing", meta["Bank Marketing"]["instances"], meta["Bank Marketing"]["encoded_features"], meta["Bank Marketing"]["clients"], meta["Bank Marketing"]["rounds"], meta["Bank Marketing"]["forget_policy"]],
            ["BAFS", "not reported", "loader-ready", "32 if present", "8 if present", "foreign/high-risk internet-origin forget requests when columns exist"],
        ],
        [0.95, 0.62, 0.78, 0.58, 0.58, 3.00],
        size=7.0,
    )
    para(doc, "Baselines. The learning-plane comparison includes SISA-Retrain, FedEraser-proxy, FedRecovery-proxy, Starfish-proxy, MA-ABE-FU, and Oracle-Retrain. Oracle-Retrain retrains on retained federated data and is used only as a reference. Because official implementations do not share a common data pipeline and model class, the package implements faithful proxies and reports a proxy-ablation table rather than hiding this limitation.")
    table_caption(doc, "IV", "Baselines and fairness boundary")
    add_table(
        doc,
        ["Method", "Family", "Implemented mechanism", "Fairness boundary"],
        [
            ["SISA-Retrain", "sharded retraining", "affected-shard retained retraining", "same features, seeds, and client partition"],
            ["FedEraser-proxy", "cached replay", "history replay with affected-client recalibration", "official same-dataset metrics unavailable"],
            ["FedRecovery-proxy", "recovery", "retained-client recovery with calibrated perturbation", "proxy for recovery behavior"],
            ["Starfish-proxy", "certified repair", "short retained-set repair proxy", "no private circuit reproduction"],
            ["MA-ABE-FU", "control plane", "same unlearning repair family plus policy capsule, padding, proofs, UCAP", "authorization contribution is isolated"],
        ],
        [1.0, 1.15, 2.15, 2.1],
        size=7.1,
    )
    para(doc, "Attacks and metrics. Membership inference follows loss-threshold attacks in which forget scope training records are members and held-out policy scope records are non-members [43], [44]. Request-type leakage measures a malicious server classifier that distinguishes forget envelopes from ordinary updates. Attribute leakage measures whether visible metadata reveals the sensitive policy attribute. Retained AUC is the primary utility metric because the banking data are imbalanced.")
    para(doc, f"Code availability. The reproducibility repository is available at {CODE_URL}. The submission package also contains exact CSV outputs, the validation script, the IEEE figure-redrawing script, dependency notes, and the BAFS status file. The BAFS loader expects files under public_data/bafs/; when no authenticated CSV is present, the runner writes a status report and skips third-dataset numeric claims.")


def add_results(doc, raw, ablation, crypto, riskgrid):
    doc.add_heading("VII. Results", level=1)
    doc.add_heading("A. Unlearning Repair Quality", level=2)
    table_caption(doc, "V", "Primary federated unlearning results averaged over seeds and forget ratios")
    add_table(
        doc,
        ["Dataset", "Method", "Ret. AUC", "MIA gap", "L2/oracle", "Runtime/oracle", "Attr. leak AUC"],
        primary_rows(raw),
        [0.95, 1.15, 0.65, 0.65, 0.72, 0.86, 0.82],
        size=7.1,
    )
    add_fig(doc, "Fig. 3.tif", "Fig. 3. Unlearning repair quality under Dirichlet non-IID alpha=0.35. The figure aggregates forget ratios 0.25, 0.50, and 1.00; error bars are 95% confidence intervals over seeds; oracle retraining is the retained-set reference.")
    para(doc, "MA-ABE-FU matches the strongest repair proxies on retained AUC while adding the policy-authenticated envelope. Averaged over datasets, seeds, and forget ratios, MA-ABE-FU obtains retained AUC 0.787, compared with 0.786 for FedRecovery-proxy and 0.786 for Starfish-proxy. Its learning-side runtime is 0.514x oracle retraining on this convex model. FedEraser-proxy remains closest to oracle in parameter distance because cached replay is especially effective for logistic training; this is why the leakage and audit results are essential rather than auxiliary.")

    doc.add_heading("B. Control-Plane Overhead and Audit Sensitivity", level=2)
    add_fig(doc, "Fig. 4.tif", "Fig. 4. Measured cryptographic overhead of the MA-ABE-FU control plane. Panel a compares primitive proxy and BN254 pairing totals on a log scale across LSSS policy rows and authority counts; panel b isolates audit-chain micro-costs.")
    para(doc, "The primitive proxy ranges from 28.2 ms at 4 policy rows to 331.7 ms at 48 rows. The BN254 pairing backend ranges from 596.0 ms to 2708.3 ms as policy rows and authorities increase. Audit-chain components remain small: HMAC, RSA-PSS verification, and hash append are minor relative to capsule/pairing work. A production implementation should therefore optimize policy compilation, batching, and pairing operations before spending effort on the hash-chain component.")
    table_caption(doc, "VI", "Measured cryptographic overhead under primitive proxy and BN254 pairing backends")
    add_table(
        doc,
        ["Backend", "Authorities", "Rows", "Capsule/pairing ms", "HMAC ms", "Sign ms", "Total ms"],
        crypto_rows(crypto),
        [1.25, 0.72, 0.55, 1.15, 0.70, 0.70, 0.75],
        size=7.0,
    )
    table_caption(doc, "VII", "Proxy calibration and RiskGap sensitivity")
    rows = proxy_rows(ablation) + [["-- RiskGap beta=2 --", "", "", "", "", ""]] + risk_rows(riskgrid)
    add_table(
        doc,
        ["Dataset/Method", "Variant/alpha", "AUC dev.", "MIA delta/RiskGap", "Runtime/pass", "Max dev./residue"],
        rows,
        [1.45, 1.05, 0.85, 1.15, 0.90, 1.05],
        size=6.8,
    )
    para(doc, "The proxy ablation supports fair comparison. FedEraser-proxy deviates from oracle by at most 0.0003 AUC on German Credit and 0.000001 on Bank Marketing when affected-client recalibration is enabled; disabling recalibration increases German Credit deviation by roughly two orders of magnitude. FedRecovery-proxy and Starfish-proxy deviate by 0.006-0.009 AUC, which bounds how much interpretation should be attached to small utility differences among repair methods.")
    para(doc, "RiskGap is best used as an audit triage score rather than a legal constant. With beta=2, alpha=0 treats privacy residue and utility loss as the primary automatic gate and passes all methods whose MIA gap is close to oracle. Alpha=0.25 exposes nonzero model distance and flags recovery-style methods for manual review. We therefore recommend a two-tier deployment: alpha=0 for automatic channel/privacy acceptance, followed by alpha=0.25 residual review whenever the authority demands model-distance accountability.")

    doc.add_heading("C. Malicious-Server Leakage", level=2)
    table_caption(doc, "VIII", "Malicious-server observability and membership residue")
    add_table(
        doc,
        ["Method", "Request-type leak AUC", "Attribute leak AUC", "MIA gap"],
        leakage_rows(raw),
        [1.45, 1.35, 1.25, 0.95],
        size=7.4,
    )
    add_fig(doc, "Fig. 5.tif", "Fig. 5. Malicious-server leakage under Dirichlet non-IID alpha=0.35 and forget ratios 0.25, 0.50, and 1.00. Points and bars report mean AUC with 95% confidence intervals; 0.5 denotes random guessing. The figure can be read independently: MA-ABE-FU lowers attribute-leakage AUC to 0.503, while FU repair proxies average 0.826.")
    para(doc, "The malicious-server leakage figure is the central privacy result. FedEraser-proxy and FedRecovery-proxy expose both request type and policy attribute almost perfectly because their observable traces are shaped by the unlearning repair scope. Starfish-proxy weakens attribute inference to 0.671 but still leaves a visible repair signature. MA-ABE-FU moves attribute leakage to 0.503, an absolute gap of 0.398 against FedEraser-proxy, 0.405 against FedRecovery-proxy, and 0.168 against Starfish-proxy. Its request-type leakage remains 0.615 rather than 0.5, which indicates that traffic padding reduces but does not completely hide timing and queue-load signals.")


def add_discussion_conclusion(doc):
    doc.add_heading("VIII. Discussion", level=1)
    para(doc, "Horizontal comparison. Federated unlearning algorithms and cryptographic authorization schemes solve complementary parts of the cross-border identity problem. FedEraser, FedRecovery, VeriFi, Starfish-style certified repair, FedWiper, and SISA-style retraining concentrate on the learning plane: their strength is utility preservation or closeness to retained-set retraining [9]-[18]. MA-ABE and decentralized ABE concentrate on the authorization plane: their strength is distributed policy enforcement [23]-[30]. MA-ABE-FU occupies the missing interface. It does not claim a better optimizer; it makes an optimizer usable when a forget request must be authorized by several authorities and hidden from a server that can profile the request surface.")
    para(doc, "Engineering tradeoffs. Four variables dominate deployment. Increasing the number of authorities improves administrative separation but raises pairing latency, as Fig. 4 shows for two, four, and six authorities. Increasing the forget ratio can improve compliance coverage but may make unlearning repair less stable under non-IID clients. Larger LSSS policies improve expressiveness but dominate control-plane latency; batching and policy-row minimization should therefore be applied before changing the audit chain. Finally, stronger model-distance gates in RiskGap increase accountability but can reject otherwise privacy-favorable repairs. A practical configuration is a two-tier gate: alpha=0 for fast privacy-channel acceptance, followed by alpha=0.25 manual review when model-distance accountability is required.")
    para(doc, "Adversarial extensions. Data poisoning and camouflaged unlearning attacks can make an apparently valid unlearning repair amplify malicious influence [16]. MA-ABE-FU can be extended with client-update provenance, robust aggregation commitments, and poison-suspicion tags inside R so that pi_rep binds not only the repair trace but also the defense configuration. Authority collusion is the second natural extension. The current game excludes enough corrupted authorities to satisfy the challenge policy. A stronger deployment should use threshold authorities, authority-diversity constraints in rho, key-issuance transparency, and periodic cross-authority audits so that no single banking, telecom, or platform authority can silently authorize a high-risk forget request.")
    table_caption(doc, "IX", "Regulatory audit mapping")
    add_table(
        doc,
        ["Regime", "Relevant provisions", "MA-ABE-FU evidence", "Boundary"],
        [
            ["GDPR", "Arts. 5(2), 17, 30, 32; SCC transfer context", "O_F records policy, scope, channel, repair commitments, proofs, and RiskGap for accountability", "does not replace legal basis assessment"],
            ["China PIPL", "Arts. 38-40 cross-border transfer; Arts. 44 and 47 individual request rights", "multi-authority policy capsule separates authorization from server-visible metadata", "authority policy must reflect local compliance review"],
            ["CCPA/CPRA", "Cal. Civ. Code Sec. 1798.105 consumer request mechanism", "UCAP chain gives timestamped evidence for request handling and audit sampling", "statutory exemptions remain outside the protocol"],
        ],
        [0.85, 1.55, 2.55, 1.35],
        size=6.9,
    )
    para(doc, "Regulatory fit. The protocol is best understood as evidence support for compliance teams, not as a legal conclusion. GDPR accountability requires the controller to demonstrate how a statutory request was handled; PIPL emphasizes individual request rights and cross-border transfer governance; CCPA/CPRA requires operational response records for consumer requests [3]-[6]. MA-ABE-FU supplies a cryptographic audit trail for who authorized the forget request, which scope was committed, whether the request channel matched the padded grammar, which unlearning repair was executed, and what residual risk remained.")
    para(doc, "Limitations and future work. The FU baselines are faithful proxies, not official end-to-end reproductions; the ablation table reports the proxy error envelope to reduce bias. The BAFS loader is present, but authenticated data were not available in the local environment, so the manuscript does not claim BAFS numeric results. The current repair model is logistic and tabular; deep biometric or multimodal identity systems would require optimizer-specific unlearning repair and stronger membership-inference tests. Future work should add public BAFS runs once authenticated data are available, implement threshold-authority collusion resistance, integrate robust aggregation into pi_rep, and benchmark production pairing libraries with hardware acceleration.")
    doc.add_heading("IX. Conclusion", level=1)
    para(doc, "MA-ABE-FU separates authorization from learning-side repair and then binds both to audit evidence. This separation matters in cross-border identity authentication, where the reason for a forget request can itself be sensitive. The experiments show that the added control plane preserves retained-set utility while reducing malicious-server attribute leakage toward random guessing. The remaining model residue is kept explicit through RiskGap, which makes the claim narrower but easier to audit.")


def add_appendix(doc):
    doc.add_heading("Appendix A. Reduction Proof of Theorem 1", level=1)
    para(doc, "Let A be a PPT adversary for Definition 1. The reduction replaces the adversary's view one component at a time. Fig. A1 gives the roadmap. The notation G_i denotes the experiment after the first i replacements.")
    appendix_proof_map(doc)
    para(doc, "G0 to G1: commitments. G0 is the real experiment. G1 samples challenge commitments through a lazy random-oracle table. Repeated queries receive the same value. A distinguisher for this hop either opens one commitment in two ways or predicts an unqueried oracle value. This yields B_2 against the random-oracle commitment abstraction, so the hop is bounded by Adv_RO(B_2)+negl(lambda).")
    para(doc, "G1 to G2: policy capsule. G2 replaces the challenge CT_F with an encryption of an independent tag under the same public LSSS shape. The adversary is not allowed enough authority keys to satisfy the challenge policy, and GID binding prevents key pooling across authorities. Any non-negligible distinction therefore gives B_1 against selective collusion-resistant MA-ABE. The hop is bounded by Adv_MA-ABE(B_1).")
    para(doc, "G2 to G3: visible channel. G3 replaces the envelope transcript by a draw from the padded update-family distribution. The two challenge events have the same public shape by construction. If A can still classify the event type, B_3 distinguishes the padding distribution. This gives the term Adv_PAD(B_3).")
    para(doc, "G3 to G4: audit chain. G4 rejects every new audit object unless it carries a valid AS signature and predecessor-chain binding. A successful forgery gives B_4 against signature unforgeability, except for commitment collisions already charged in G1. The additional loss is bounded by Adv_SIG(B_4).")
    para(doc, "G4 to G5: proof soundness. G5 rejects if pi_auth or pi_rep accepts a false statement. A false accepted authorization or repair proof gives B_5 against knowledge soundness. This hop is bounded by Adv_ZK(B_5).")
    para(doc, "Final game. In G5 the adversary observes padded metadata, random commitments, valid evidence objects, and the empirical residual report. The only remaining distinguishing signal is the configured residual-risk tolerance eps_R for the selected learning module. Summing the five hops proves Theorem 1.")


def add_references(doc):
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFS, 1):
        para(doc, f"[{i}] {ref}", size=8.7)


def build_manuscript():
    raw = load_csv("federated_raw_v6.csv")
    ablation = load_csv("proxy_ablation_v6.csv")
    crypto = load_csv("crypto_overhead_v6.csv")
    riskgrid = load_csv("riskgap_sensitivity_v6.csv")
    meta = json.loads((REPRO / "validation_metadata_v6.json").read_text(encoding="utf-8"))
    bafs_status = json.loads((REPRO / "bafs_status_v6.json").read_text(encoding="utf-8"))

    doc = Document()
    setup_doc(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15.0)
    para(doc, AUTHORS, center=True, size=10.8)
    doc.add_heading("Abstract", level=1)
    para(doc, "Federated unlearning is often treated as an unlearning repair task. A server receives a forget request and then reports how closely the repaired model matches retained-set retraining. Cross-border identity authentication adds a separate control problem. The request may need approval from several authorities, yet the server should not learn sensitive policy attributes from request metadata. This paper presents MA-ABE-FU, a control plane that uses multi-authority attribute-based encryption to authorize the request channel and uses audit evidence to bind the resulting repair. Its security game separates request hiding and audit binding from the empirical residue left by the learning algorithm. Experiments on non-IID German Credit and Bank Marketing partitions evaluate forget ratios, membership inference, malicious-server request classification, attribute leakage, proxy baselines, and measured cryptographic overhead. MA-ABE-FU keeps retained AUC close to representative repair baselines and lowers mean attribute-leakage AUC to 0.503. The result is not a claim that encryption makes model influence vanish. It is an auditable way to run authorized federated forget requests without exposing the policy surface to the server.")
    para(doc, "Index Terms - Federated unlearning, multi-authority attribute-based encryption, cross-border identity authentication, zero-knowledge proof, membership inference, audit evidence.")
    start_two_column_body(doc)
    add_introduction(doc, raw)
    add_related_work(doc)
    add_threat_model(doc)
    add_protocol(doc)
    add_security(doc)
    add_experiments(doc, meta, bafs_status)
    add_results(doc, raw, ablation, crypto, riskgrid)
    add_discussion_conclusion(doc)
    add_appendix(doc)
    add_references(doc)
    doc.save(OUT / "manuscript.docx")


def simple_doc(filename: str, lines):
    doc = Document()
    setup_doc(doc)
    for i, line in enumerate(lines):
        para(doc, line, bold=(i == 0), center=(i == 0), size=12 if i == 0 else 10.0)
    doc.save(OUT / filename)


def other_docs():
    simple_doc(
        "Title page.docx",
        [
            TITLE,
            "Jian Chen a,b, Sheng Peng b, Zhiming Cai b,c, Jiayin Qi d, Fu Mo e,*",
            "a Faculty of Finance and Economics, Guangdong University of Science and Technology, Dongguan, 523083, China",
            "b Faculty of Digital Science and Technology, Macau Millennium College, Macau, 999078, China",
            "c Institute for Data Engineering and Science, University of Saint Joseph, Macao, 999078, China",
            "d Cyberspace Institute of Advanced Technology, Guangzhou University, Guangzhou, 510000, China",
            "e Institute of Data Intelligence, Guangdong University of Science and Technology, Dongguan, 523083, China",
            f"*Corresponding author. Email: {EMAIL}. Tel: {TEL}.",
            "Funding",
            "This work was supported by the Macao FDCT (No. 002/2025/STT); the Chinese Academy of Engineering (No. 2025-XZ-08); and the 2025 Guangdong Provincial Project for Enhancing Research Capacity of Key Construction Disciplines (No. 2025ZDJS068).",
            "Declaration of competing interest",
            "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this work.",
            "Author contribution statement",
            "Jian Chen: Conceptualization, Methodology, Writing - Original Draft. Sheng Peng: Data Curation, Formal Analysis. Zhiming Cai: Validation, Formal Analysis. Jiayin Qi: Supervision, Project Administration, Writing - Review and Editing. Fu Mo: Supervision, Funding Acquisition, Resources, Writing - Review and Editing.",
        ],
    )
    simple_doc(
        "Cover Letter.docx",
        [
            "Cover Letter",
            "July 19, 2026",
            f"Editor-in-Chief, {TARGET}",
            "Dear Editor-in-Chief,",
            f"We submit \"{TITLE}\" as a regular research article.",
            "The manuscript presents a policy-authenticated and evidence-bound control plane for federated unlearning in cross-border identity authentication. It includes a complete protocol, formal security game, reduction proof, non-IID federated experiments, malicious-server observability tests, proxy-baseline ablations, and measured cryptographic overhead.",
            f"The code and reproducibility material are available at {CODE_URL}.",
            "The submission fits the scope of IEEE Transactions on Information Forensics and Security because it addresses privacy leakage, applied cryptographic authorization, zero-knowledge auditability, and evidence-bound federated unlearning for identity authentication.",
            "The manuscript is original, not under consideration elsewhere, and approved by all authors.",
            "Sincerely,",
            "Prof. Fu Mo",
            f"Email: {EMAIL}; Tel: {TEL}",
        ],
    )
    simple_doc(
        "Highlights.docx",
        [
            "Highlights",
            "MA-ABE-FU provides a policy-authenticated control plane for federated unlearning.",
            "The protocol defines four algorithms, two proof relations, and one formal update-hiding game.",
            "The reduction proof separates cryptographic guarantees from empirical model residue.",
            "Experiments cover non-IID clients, forget ratios, leakage attacks, and FU proxy ablations.",
            "The package reports measured primitive and BN254 pairing control-plane overhead.",
        ],
    )
    final_bio = next((p for p in (ROOT.parent / "final version").glob("Author*biography.docx")), None)
    if final_bio and final_bio.exists():
        shutil.copy2(final_bio, OUT / "Author biography.docx")
    else:
        simple_doc("Author biography.docx", ["Author Biography"] + BIOS)


def supplement():
    raw = load_csv("federated_raw_v6.csv")
    summary = load_csv("federated_summary_v6.csv")
    ablation = load_csv("proxy_ablation_v6.csv")
    crypto = load_csv("crypto_overhead_v6.csv")
    riskgrid = load_csv("riskgap_sensitivity_v6.csv")
    bafs_status = json.loads((REPRO / "bafs_status_v6.json").read_text(encoding="utf-8"))
    doc = Document()
    setup_doc(doc)
    para(doc, "Supplementary Material", bold=True, center=True, size=14)
    doc.add_heading("A. Artifact Manifest", level=1)
    para(doc, f"The reproducibility repository is {CODE_URL}. The submission archive contains run_validation_v6.py, redraw_ieee_figures_v6.py, prebuild_validation_v6.py, federated_raw_v6.csv, federated_summary_v6.csv, proxy_ablation_v6.csv, riskgap_sensitivity_v6.csv, crypto_overhead_v6.csv, validation_metadata_v6.json, term_audit_v6.csv, symbol_audit_v6.csv, figure_audit_v6.csv, format_audit_v6.csv, reference_audit_v6.csv, and bafs_status_v6.json.")
    para(doc, f"BAFS status: available={bafs_status['available']}; expected files: {', '.join(bafs_status['expected_files'])}. No BAFS numeric row is used unless a listed CSV is present.")
    table_caption(doc, "S-I", "Federated result summary by ratio")
    add_table(
        doc,
        ["Dataset", "Ratio", "Method", "Ret. AUC", "MIA gap", "L2", "Type leak", "Attr leak"],
        [
            [r["dataset"], r["forget_ratio"], r["method"], fnum(r["retained_auc_mean"], 3), fnum(r["mia_gap_mean"], 3), fnum(r["l2_to_oracle_mean"], 3), fnum(r["type_leak_auc_mean"], 3), fnum(r["attribute_leak_auc_mean"], 3)]
            for r in summary
            if r["method"] in ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"]
        ],
        [0.85, 0.45, 1.05, 0.62, 0.58, 0.55, 0.66, 0.66],
        size=6.5,
    )
    table_caption(doc, "S-II", "Proxy ablation cells")
    add_table(
        doc,
        ["Dataset", "Method", "Ablation", "AUC dev.", "MIA delta", "Runtime/oracle", "Max AUC dev."],
        [[r["dataset"], r["method"], r["ablation"], fnum(r["mean_abs_auc_delta_to_oracle"], 5), fnum(r["mean_mia_gap_delta_to_oracle"], 5), fnum(r["mean_runtime_to_oracle"], 3), fnum(r["max_abs_auc_delta_to_oracle"], 5)] for r in ablation],
        [0.75, 1.05, 1.35, 0.75, 0.75, 0.83, 0.83],
        size=6.4,
    )
    table_caption(doc, "S-III", "RiskGap grid excerpt")
    add_table(
        doc,
        ["alpha", "beta", "Method", "RiskGap", "Pass", "Residue", "Privacy gap", "Utility loss"],
        [[r["alpha"], r["beta"], r["method"], fnum(r["mean_riskgap"], 3), fnum(r["audit_pass_rate"], 2), fnum(r["mean_model_residue"], 3), fnum(r["mean_privacy_gap"], 3), fnum(r["mean_utility_loss"], 4)] for r in riskgrid if float(r["beta"]) in {0.0, 2.0} and float(r["alpha"]) in {0.0, 0.25, 1.0}],
        [0.5, 0.45, 1.05, 0.65, 0.45, 0.65, 0.68, 0.68],
        size=6.2,
    )
    table_caption(doc, "S-IV", "Cryptographic benchmark full table")
    add_table(
        doc,
        ["Backend", "Auth.", "Rows", "Capsule", "Pairing", "HMAC", "Sign", "Verify", "Hash", "Total"],
        [[r["backend"], r["authority_count"], r["policy_rows"], fnum(r["ma_abe_capsule_proxy_ms"], 2) if r["ma_abe_capsule_proxy_ms"] else "", fnum(r["bn254_pairing_ma_abe_ms"], 2) if r["bn254_pairing_ma_abe_ms"] else "", fnum(r["padded_envelope_hmac_ms"], 3), fnum(r["ucap_rsa_pss_sign_ms"], 3), fnum(r["ucap_rsa_pss_verify_ms"], 3), fnum(r["hash_chain_append_ms"], 3), fnum(r["total_control_plane_ms"], 2)] for r in crypto],
        [1.10, 0.42, 0.42, 0.63, 0.63, 0.52, 0.52, 0.52, 0.52, 0.63],
        size=5.9,
    )
    para(doc, f"Raw run count: {len(raw)} federated records. The CSV files in the reproducibility folder are the authoritative source for all plotted values.")
    doc.save(OUT / "Supplementary material.docx")


def bib_escape(text: str) -> str:
    return text.replace("\\", "\\textbackslash{}").replace("{", "\\{").replace("}", "\\}")


def write_references_bib():
    entries = []
    for i, ref in enumerate(REFS, 1):
        entries.append(f"@misc{{ref{i:02d},\n  note = {{{bib_escape(ref)}}}\n}}")
    (OUT / "references.bib").write_text("\n\n".join(entries) + "\n", encoding="utf-8")


def copy_reproducibility_files():
    if REPRO.exists():
        stale_names = {
            "federated_raw_v4.csv",
            "federated_summary_v4.csv",
            "crypto_overhead_v4.csv",
            "validation_metadata_v4.json",
            "run_validation_v4.py",
            "redraw_ieee_figures.py",
        }
        for item in REPRO.iterdir():
            if item.is_file() and (item.name in stale_names or "_v4" in item.name):
                item.unlink()
    shutil.copy2(ROOT / "run_validation_v6.py", REPRO / "run_validation_v6.py")
    shutil.copy2(ROOT / "redraw_ieee_figures_v6.py", REPRO / "redraw_ieee_figures_v6.py")
    shutil.copy2(ROOT / "prebuild_validation_v6.py", REPRO / "prebuild_validation_v6.py")
    README = f"""# MA-ABE-FU reproducibility package

Target manuscript: {TITLE}

Public repository: {CODE_URL}

## Contents

- `run_validation_v6.py`: federated partitioning, learning-plane baselines, leakage attacks, proxy ablations, RiskGap sensitivity, and cryptographic benchmarks.
- `redraw_ieee_figures_v6.py`: 600 dpi IEEE-style figure generation from CSV results.
- `prebuild_validation_v6.py`: terminology, symbol, figure, format, reference, and trace audit before manuscript construction.
- `build_submission_v6.py`: manuscript, supplementary material, title page, highlights, cover letter, and package builder.
- `submission_tifs_v6/reproducibility/*.csv|*.json`: exact experiment outputs and validation records used in the manuscript.
- `submission_tifs_v6/figure/Fig. *.pdf|*.tif`: vector PDFs and 600 dpi TIFF figures.
- `references.bib`: IEEE-formatted reference list exported for submission support.

## Reproduction

Install Python dependencies:

```bash
python -m pip install numpy pandas pillow cryptography py_ecc python-docx
```

Install Poppler, or set `PDFTOCAIRO` to the absolute path of `pdftocairo`, before redrawing figures.

Place public datasets under `public_data/`:

- `public_data/german.data`
- `public_data/bank/bank-full.csv`
- optional BAFS CSV files under `public_data/bafs/`

Run:

```bash
python run_validation_v6.py
python redraw_ieee_figures_v6.py
python prebuild_validation_v6.py
python build_submission_v6.py
```

The scripts regenerate all manuscript tables, exactly five main figures as vector PDFs plus 600 dpi TIFF files, the validation records, and the manuscript package. If BAFS files are absent, the runner writes `bafs_status_v6.json` and skips BAFS numeric results.
"""
    (OUT / "README_submission.md").write_text(README, encoding="utf-8")


def build_local_repro_repo():
    if REPRO_REPO.exists():
        shutil.rmtree(REPRO_REPO)
    (REPRO_REPO / "results").mkdir(parents=True)
    (REPRO_REPO / "figures").mkdir(parents=True)
    (REPRO_REPO / "submission_tifs_v6" / "reproducibility").mkdir(parents=True)
    (REPRO_REPO / "submission_tifs_v6" / "figure").mkdir(parents=True)
    shutil.copy2(ROOT / "run_validation_v6.py", REPRO_REPO / "run_validation_v6.py")
    shutil.copy2(ROOT / "redraw_ieee_figures_v6.py", REPRO_REPO / "redraw_ieee_figures_v6.py")
    shutil.copy2(ROOT / "prebuild_validation_v6.py", REPRO_REPO / "prebuild_validation_v6.py")
    shutil.copy2(ROOT / "build_submission_v6.py", REPRO_REPO / "build_submission_v6.py")
    shutil.copy2(OUT / "references.bib", REPRO_REPO / "references.bib")
    for name in [
        "federated_raw_v6.csv",
        "federated_summary_v6.csv",
        "proxy_ablation_v6.csv",
        "riskgap_sensitivity_v6.csv",
        "crypto_overhead_v6.csv",
        "validation_metadata_v6.json",
        "bafs_status_v6.json",
        "figure_manifest_v6.csv",
        "term_audit_v6.csv",
        "symbol_audit_v6.csv",
        "figure_audit_v6.csv",
        "format_audit_v6.csv",
        "reference_audit_v6.csv",
        "trace_audit_v6.csv",
        "prebuild_validation_v6.json",
    ]:
        src = REPRO / name
        if src.exists():
            shutil.copy2(src, REPRO_REPO / "results" / name)
            shutil.copy2(src, REPRO_REPO / "submission_tifs_v6" / "reproducibility" / name)
    for fig in sorted(FIG.glob("Fig. *.*")):
        shutil.copy2(fig, REPRO_REPO / "figures" / fig.name)
        shutil.copy2(fig, REPRO_REPO / "submission_tifs_v6" / "figure" / fig.name)
    (REPRO_REPO / "requirements.txt").write_text("numpy\npandas\npillow\ncryptography\npy_ecc\npython-docx\n", encoding="utf-8")
    (REPRO_REPO / "README.md").write_text((OUT / "README_submission.md").read_text(encoding="utf-8"), encoding="utf-8")
    (REPRO_REPO / "DATASETS.md").write_text(
        """# Dataset placement

The experiments use UCI German Credit and UCI Bank Marketing. The submission archive contains the local copies used for the reported run. For a public repository, place the same files as follows:

- `public_data/german.data`
- `public_data/bank/bank-full.csv`

The optional Bank Account Fraud Dataset Suite is evaluated only when one of these files is present:

- `public_data/bafs/Base.csv`
- `public_data/bafs/Variant I.csv`
- `public_data/bafs/Variant II.csv`
- `public_data/bafs/Variant III.csv`
- `public_data/bafs/Variant IV.csv`
- `public_data/bafs/Variant V.csv`

If no BAFS CSV is present, `run_validation_v6.py` writes `bafs_status_v6.json` and does not generate third-dataset metrics.
""",
        encoding="utf-8",
    )
    (REPRO_REPO / "PUBLISH.md").write_text(
        """# Repository publication

The directory is ready to publish as the reproducibility repository for the manuscript.

Before using SSH, add the local public key to GitHub:

1. Open GitHub Settings.
2. Go to SSH and GPG keys.
3. Add the content of the local public key file, typically `~/.ssh/id_ed25519.pub`, as a new SSH key.
4. Verify with `ssh -T git@github.com`.

```bash
git init -b main
git add .
git commit -m "Add MA-ABE-FU reproducibility package"
git remote add origin https://github.com/Kent919/MA-ABE-FU.git
git push -u origin main
```

If the remote already exists, replace the `git remote add` command with:

```bash
git remote set-url origin https://github.com/Kent919/MA-ABE-FU.git
```

The repository contains scripts, exact CSV/JSON outputs, five vector PDF figures, five 600 dpi TIFF figures, and the reference export. Public datasets should remain outside version control unless redistribution rights are clear.
""",
        encoding="utf-8",
    )
    (REPRO_REPO / "push_to_github.sh").write_text(
        """#!/usr/bin/env bash
set -euo pipefail

git remote set-url origin git@github.com:Kent919/MA-ABE-FU.git
git push -u origin main
""",
        encoding="utf-8",
    )


def package_zip():
    zip_path = ROOT / "MA-ABE-FU_TIFS_submission_package_v6.zip"
    if zip_path.exists():
        zip_path.unlink()
    def should_package(path: Path) -> bool:
        parts = set(path.parts)
        if path.name == ".DS_Store" or path.name.startswith(".~"):
            return False
        if "__pycache__" in parts or "render_check" in parts:
            return False
        return True
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
        for f in FINAL_LAYOUT.rglob("*"):
            if f.is_file() and should_package(f):
                z.write(f, f.relative_to(FINAL_LAYOUT))
    return zip_path


def copy_tree_filtered(src: Path, dst: Path) -> None:
    if not src.exists():
        return
    for path in src.rglob("*"):
        if path.is_file() and path.name != ".DS_Store" and not path.name.startswith(".~"):
            target = dst / path.relative_to(src)
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(path, target)


def prepare_final_layout() -> None:
    if FINAL_LAYOUT.exists():
        shutil.rmtree(FINAL_LAYOUT)
    (SUBMISSION_REQUIRED / "figure").mkdir(parents=True)
    (SUPPORTING_FILES / "reproducibility_repository").mkdir(parents=True)
    required_files = [
        "manuscript.docx",
        "manuscript.pdf",
        "Title page.docx",
        "Cover Letter.docx",
        "Highlights.docx",
        "Author biography.docx",
        "Supplementary material.docx",
        "references.bib",
        "README_submission.md",
    ]
    for name in required_files:
        src = OUT / name
        if src.exists():
            shutil.copy2(src, SUBMISSION_REQUIRED / name)
    for fig in sorted(FIG.glob("Fig. *.*")):
        shutil.copy2(fig, SUBMISSION_REQUIRED / "figure" / fig.name)
    copy_tree_filtered(REPRO_REPO, SUPPORTING_FILES / "reproducibility_repository")
    audit_dir = SUPPORTING_FILES / "audits"
    audit_dir.mkdir(parents=True, exist_ok=True)
    for audit in REPRO.glob("*audit*_v6.csv"):
        shutil.copy2(audit, audit_dir / audit.name)
    for meta in ["prebuild_validation_v6.json", "figure_manifest_v6.csv", "bafs_status_v6.json"]:
        src = REPRO / meta
        if src.exists():
            shutil.copy2(src, audit_dir / meta)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    FIG.mkdir(parents=True, exist_ok=True)
    REPRO.mkdir(parents=True, exist_ok=True)
    required = [
        REPRO / "federated_raw_v6.csv",
        REPRO / "federated_summary_v6.csv",
        REPRO / "proxy_ablation_v6.csv",
        REPRO / "riskgap_sensitivity_v6.csv",
        REPRO / "crypto_overhead_v6.csv",
        REPRO / "validation_metadata_v6.json",
        REPRO / "bafs_status_v6.json",
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("Run run_validation_v6.py before building the submission package: " + ", ".join(missing))
    if len(list(FIG.glob("Fig. *.tif"))) != 5:
        raise RuntimeError("Run redraw_ieee_figures_v6.py before building; exactly five main TIFF figures are required.")
    copy_reproducibility_files()
    build_manuscript()
    other_docs()
    supplement()
    write_references_bib()
    build_local_repro_repo()
    prepare_final_layout()
    zip_path = package_zip()
    print(zip_path)


if __name__ == "__main__":
    main()
