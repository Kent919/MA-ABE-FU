#!/usr/bin/env python3
"""Build the MA-ABE-FU v5 IEEE submission package."""

from __future__ import annotations

import csv
import json
import re
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "submission_tdsc_v5"
FIG = OUT / "figure"
REPRO = OUT / "reproducibility"
PUBLIC_DATA = ROOT / "public_data"
REPRO_REPO = ROOT / "ma-abe-fu-reproducibility"

TITLE = "MA-ABE-FU: Policy-Authenticated and Evidence-Bound Federated Unlearning for Cross-Border Identity Authentication"
AUTHORS = "Jian Chen, Sheng Peng, Zhiming Cai, Jiayin Qi, and Fu Mo"
EMAIL = "mofu@gdust.edu.cn"
TEL = "+86 15989698699"
TARGET = "IEEE Transactions on Dependable and Secure Computing"
CODE_URL = "https://github.com/Kent919/MA-ABE-FU"

REFS = [
    "Financial Action Task Force, \"Guidance on Digital Identity,\" Paris, France, 2020.",
    "European Banking Authority, \"Guidelines on the use of remote customer onboarding solutions,\" EBA/GL/2022/15, 2022.",
    "European Parliament and Council of the European Union, Regulation (EU) 2016/679, General Data Protection Regulation, arts. 5(2), 17, 30, and 32, 2016.",
    "European Commission, Commission Implementing Decision (EU) 2021/914 on standard contractual clauses for the transfer of personal data to third countries, clauses 8 and 15, 2021.",
    "Standing Committee of the National People's Congress, Personal Information Protection Law of the People's Republic of China, arts. 38-40, 44, and 47, 2021.",
    "California State Legislature, California Consumer Privacy Act of 2018, Cal. Civ. Code Sec. 1798.105, 2018.",
    "B. McMahan, E. Moore, D. Ramage, S. Hampson, and B. A. y Arcas, \"Communication-Efficient Learning of Deep Networks From Decentralized Data,\" in Proc. AISTATS, pp. 1273-1282, 2017.",
    "P. Kairouz et al., \"Advances and Open Problems in Federated Learning,\" Found. Trends Mach. Learn., vol. 14, no. 1-2, pp. 1-210, 2021, doi: 10.1561/2200000083.",
    "G. Liu, X. Ma, Y. Yang, C. Wang, and J. Liu, \"FedEraser: Enabling Efficient Client-Level Data Removal From Federated Learning Models,\" in Proc. IEEE/ACM IWQoS, pp. 1-10, 2021, doi: 10.1109/IWQOS52092.2021.9521274.",
    "L. Zhang, T. Zhu, H. Zhang, P. Xiong, and W. Zhou, \"FedRecovery: Differentially Private Machine Unlearning for Federated Learning Frameworks,\" IEEE Trans. Inf. Forensics Secur., vol. 18, pp. 4732-4746, 2023, doi: 10.1109/TIFS.2023.3297905.",
    "X. Gao, X. Ma, J. Wang, Y. Sun, B. Li, S. Ji, P. Cheng, and J. Chen, \"VeriFi: Towards Verifiable Federated Unlearning,\" IEEE Trans. Dependable Secure Comput., vol. 21, no. 6, pp. 5720-5736, Nov. 2024, doi: 10.1109/TDSC.2024.3382321.",
    "Z. Liu, Y. Jiang, W. Jiang, J. Guo, J. Zhao, and K.-Y. Lam, \"Guaranteeing Data Privacy in Federated Unlearning With Dynamic User Participation,\" IEEE Trans. Dependable Secure Comput., vol. 22, no. 3, pp. 2072-2085, May 2025, doi: 10.1109/TDSC.2024.3476533.",
    "J. Chen, Z. Lin, W. Lin, W. Shi, X. Yin, and D. Wang, \"FedMUA: Exploring the Vulnerabilities of Federated Learning to Malicious Unlearning Attacks,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 1665-1678, 2025, doi: 10.1109/TIFS.2025.3531141.",
    "Z. Liu, H. Ye, Y. Jiang, J. Shen, J. Guo, I. Tjuawinata, and K.-Y. Lam, \"Privacy-Preserving Federated Unlearning With Certified Client Removal,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 3966-3978, 2025, doi: 10.1109/TIFS.2025.3555868.",
    "S. Zhao, J. Zhang, X. Ma, Q. Jiang, Z. Ma, S. Gao, Z. Ying, and J. Ma, \"FedWiper: Federated Unlearning via Universal Adapter,\" IEEE Trans. Inf. Forensics Secur., vol. 20, pp. 4042-4054, 2025, doi: 10.1109/TIFS.2025.3557671.",
    "K. Gao, T. Zhu, D. Ye, B. Liu, and W. Zhou, \"Hidden Threats in Federated Unlearning: Camouflaged Poisoning Attacks and Their Unlearning Consequences,\" IEEE Trans. Dependable Secure Comput., vol. 23, no. 2, pp. 2934-2948, Mar. 2026, doi: 10.1109/TDSC.2025.3630811.",
    "L. Bourtoule et al., \"Machine Unlearning,\" in Proc. IEEE Symp. Security and Privacy, pp. 141-159, 2021, doi: 10.1109/SP40001.2021.00019.",
    "C. Guo, T. Goldstein, A. Hannun, and L. van der Maaten, \"Certified Data Removal From Machine Learning Models,\" in Proc. ICML, pp. 3832-3842, 2020.",
    "S. Garg, S. Goldwasser, and P. N. Vasudevan, \"Formalizing Data Deletion in the Context of the Right to Be Forgotten,\" in EUROCRYPT, pp. 373-402, 2020.",
    "Z. Liu, Y. Jiang, J. Shen, M. Peng, K.-Y. Lam, X. Yuan, and X. Liu, \"A Survey on Federated Unlearning: Challenges, Methods, and Future Directions,\" ACM Comput. Surv., vol. 57, no. 1, pp. 1-38, Jan. 2025, doi: 10.1145/3679014.",
    "T. T. Nguyen et al., \"A Survey of Machine Unlearning,\" ACM Trans. Intell. Syst. Technol., vol. 16, no. 5, Art. 108, 2025, doi: 10.1145/3749987.",
    "F. Wang, B. Li, and B. Li, \"Federated Unlearning and Its Privacy Threats,\" IEEE Network, vol. 38, no. 2, pp. 294-300, Mar.-Apr. 2024, doi: 10.1109/MNET.004.2300056.",
    "A. Sahai and B. Waters, \"Fuzzy Identity-Based Encryption,\" in EUROCRYPT, pp. 457-473, 2005.",
    "V. Goyal, O. Pandey, A. Sahai, and B. Waters, \"Attribute-Based Encryption for Fine-Grained Access Control of Encrypted Data,\" in Proc. ACM CCS, pp. 89-98, 2006.",
    "J. Bethencourt, A. Sahai, and B. Waters, \"Ciphertext-Policy Attribute-Based Encryption,\" in Proc. IEEE Symp. Security and Privacy, pp. 321-334, 2007.",
    "M. Chase, \"Multi-Authority Attribute Based Encryption,\" in TCC, pp. 515-534, 2007.",
    "A. Lewko and B. Waters, \"Decentralizing Attribute-Based Encryption,\" in EUROCRYPT, pp. 568-588, 2011.",
    "B. Waters, \"Ciphertext-Policy Attribute-Based Encryption: An Expressive, Efficient, and Provably Secure Realization,\" in PKC, pp. 53-70, 2011.",
    "Y. Zhang, R. H. Deng, S. Xu, J. Sun, Q. Li, and D. Wu, \"Attribute-Based Encryption for Cloud Computing Access Control: A Survey,\" ACM Comput. Surv., vol. 53, no. 4, pp. 1-37, 2020.",
    "X. Xing, Y. Liu, Q. Wu, Z. Guan, D. Li, D. Li, Y. Lu, and W. Susilo, \"Multi-Committee ABE Based Decentralized Access Control With Sharding Blockchain for Web 3.0,\" IEEE Trans. Dependable Secure Comput., vol. 22, no. 3, pp. 2533-2549, May 2025, doi: 10.1109/TDSC.2024.3520121.",
    "D. Boneh and M. Franklin, \"Identity-Based Encryption From the Weil Pairing,\" in CRYPTO, pp. 213-229, 2001.",
    "A. Fiat and A. Shamir, \"How to Prove Yourself: Practical Solutions to Identification Problems,\" in CRYPTO, pp. 186-194, 1986.",
    "S. Goldwasser, S. Micali, and C. Rackoff, \"The Knowledge Complexity of Interactive Proof Systems,\" SIAM J. Comput., vol. 18, no. 1, pp. 186-208, 1989.",
    "J. Groth, \"On the Size of Pairing-Based Non-Interactive Arguments,\" in EUROCRYPT, pp. 305-326, 2016.",
    "B. Bunz et al., \"Bulletproofs: Short Proofs for Confidential Transactions and More,\" in Proc. IEEE Symp. Security and Privacy, pp. 315-334, 2018.",
    "A. Gabizon, Z. J. Williamson, and O. Ciobotaru, \"PLONK: Permutations Over Lagrange-Bases for Oecumenical Noninteractive Arguments of Knowledge,\" Cryptology ePrint Archive, 2019/953, 2019.",
    "E. Ben-Sasson, I. Bentov, Y. Horesh, and M. Riabzev, \"Scalable, Transparent, and Post-Quantum Secure Computational Integrity,\" Cryptology ePrint Archive, 2018/046, 2018.",
    "E. Androulaki et al., \"Hyperledger Fabric: A Distributed Operating System for Permissioned Blockchains,\" in Proc. EuroSys, pp. 1-15, 2018.",
    "S. Nakamoto, \"Bitcoin: A Peer-to-Peer Electronic Cash System,\" 2008.",
    "H. Hofmann, \"Statlog (German Credit Data),\" UCI Machine Learning Repository, 1994, doi: 10.24432/C5NC77.",
    "S. Moro, P. Rita, and P. Cortez, \"Bank Marketing,\" UCI Machine Learning Repository, 2014, doi: 10.24432/C5K306.",
    "S. Jesus, J. Pombal, D. Alves, A. Cruz, P. Saleiro, R. Ribeiro, J. Gama, and P. Bizarro, \"Turning the Tables: Biased, Imbalanced, Dynamic Tabular Datasets for ML Evaluation,\" in Adv. Neural Inf. Process. Syst. 35, pp. 33563-33575, 2022, doi: 10.52202/068431-2432.",
    "R. Shokri, M. Stronati, C. Song, and V. Shmatikov, \"Membership Inference Attacks Against Machine Learning Models,\" in Proc. IEEE Symp. Security and Privacy, pp. 3-18, 2017.",
    "S. Yeom, I. Giacomelli, M. Fredrikson, and S. Jha, \"Privacy Risk in Machine Learning: Analyzing the Connection to Overfitting,\" in Proc. IEEE CSF, pp. 268-282, 2018.",
]

BIOS = [
    "Jian Chen, PhD. is a senior researcher at Macau Millennium College and a member of the Macau-Qin'ao Cross-Border Data Flow Research Team. His research interests include smart cities and digital security.",
    "Jiayin Qi, PhD. is a professor of Cyberspace Security at Guangzhou University and executive director of the Institute for Innovation in Cyberspace Governance, Guangdong-Hong Kong-Macao Greater Bay Area. Her research interests include cross-border data governance, privacy-preserving computation, and trustworthy data sharing.",
    "Sheng Peng, PhD. is an associate professor at Macau Millennium College and a postdoctoral research fellow. His research interests include cross-border digital identity, large-scale system implementation, and data security engineering.",
    "Zhiming Cai, PhD. is a professor at Macau Millennium College. His research interests include cross-border digital identity, large-scale system implementation, and data security engineering.",
    "Fu Mo. PhD. is an associate professor at The Institute of Data Intelligence, Guangdong University of Science and Technology, where he also serves as deputy director of the research affairs office. His research interests include intelligent electronic systems and data analytics, with particular focus on big-data analysis and decision-support technologies.",
]


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.70)
    sec.right_margin = Inches(0.70)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(10.0)
    doc.styles["Heading 1"].font.size = Pt(12.0)
    doc.styles["Heading 2"].font.size = Pt(11.0)
    doc.styles["Heading 3"].font.size = Pt(10.5)


def para(doc: Document, text: str = "", bold: bool = False, italic: bool = False, center: bool = False, size: float = 10.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.0)
    p.paragraph_format.line_spacing = 1.04
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    return p


def formula(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(9.3)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def looks_numeric(value) -> bool:
    try:
        float(str(value).strip())
        return True
    except ValueError:
        return False


def cell_text(cell, value, bold: bool = False, size: float = 8.0, align=None) -> None:
    cell.text = ""
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    if align is None:
        align = WD_ALIGN_PARAGRAPH.CENTER if looks_numeric(value) else WD_ALIGN_PARAGRAPH.LEFT
    p.alignment = align
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_caption(doc: Document, num: str, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"TABLE {num}\n{title.upper()}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)


def add_table(doc: Document, headers, rows, widths=None, size: float = 7.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(t.rows[0].cells[i], "EAF0F7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, False, size=size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    para(doc)
    return t


def add_fig(doc: Document, filename: str, caption: str, width: float = 6.92):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(7)
    r = c.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)


def alg_box(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, "F6F8FB")
    set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(8.8)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rr = p2.add_run(body)
    rr.font.name = "Courier New"
    rr._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    rr.font.size = Pt(7.2)
    para(doc)


def load_csv(name: str):
    with open(REPRO / name, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fnum(value, digits=3) -> str:
    return f"{float(value):.{digits}f}"


def mean(rows, filters, col: str) -> float:
    vals = [float(r[col]) for r in rows if all(r[k] == v for k, v in filters.items())]
    return sum(vals) / len(vals)


def primary_rows(raw):
    methods = ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"]
    rows = []
    for dataset in ["German Credit", "Bank Marketing"]:
        for method in methods:
            filt = {"dataset": dataset, "method": method}
            rows.append(
                [
                    dataset,
                    method,
                    fnum(mean(raw, filt, "retained_auc"), 3),
                    fnum(mean(raw, filt, "mia_gap"), 3),
                    fnum(mean(raw, filt, "l2_to_oracle"), 3),
                    fnum(mean(raw, filt, "runtime_to_oracle"), 3),
                    fnum(mean(raw, filt, "attribute_leak_auc"), 3),
                ]
            )
    return rows


def leakage_rows(raw):
    methods = ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"]
    rows = []
    for method in methods:
        filt = {"method": method}
        rows.append(
            [
                method,
                fnum(mean(raw, filt, "type_leak_auc"), 3),
                fnum(mean(raw, filt, "attribute_leak_auc"), 3),
                fnum(mean(raw, filt, "mia_gap"), 3),
            ]
        )
    return rows


def proxy_rows(ablation):
    rows = []
    for r in ablation:
        if r["ablation"] != "full proxy":
            continue
        rows.append(
            [
                r["dataset"],
                r["method"].replace("-proxy", ""),
                fnum(r["mean_abs_auc_delta_to_oracle"], 4),
                fnum(r["mean_mia_gap_delta_to_oracle"], 4),
                fnum(r["mean_runtime_to_oracle"], 3),
                fnum(r["max_abs_auc_delta_to_oracle"], 4),
            ]
        )
    return rows


def crypto_rows(crypto):
    rows = []
    for r in crypto:
        backend = r["backend"]
        policy_rows = int(float(r["policy_rows"]))
        auth = int(float(r["authority_count"]))
        keep = backend == "primitive_modexp_proxy" and policy_rows in {4, 16, 32, 48}
        keep = keep or (backend == "bn254_pairing_py_ecc" and policy_rows in {4, 24, 48} and auth in {2, 4, 6})
        if not keep:
            continue
        capsule = r["ma_abe_capsule_proxy_ms"] if backend == "primitive_modexp_proxy" else r["bn254_pairing_ma_abe_ms"]
        rows.append(
            [
                "primitive proxy" if backend == "primitive_modexp_proxy" else "BN254 pairing",
                "-" if auth == 0 else str(auth),
                str(policy_rows),
                fnum(capsule, 1),
                fnum(r["padded_envelope_hmac_ms"], 3),
                fnum(r["ucap_rsa_pss_sign_ms"], 3),
                fnum(r["total_control_plane_ms"], 1),
            ]
        )
    return rows


def risk_rows(riskgrid):
    rows = []
    for r in riskgrid:
        if float(r["beta"]) == 2.0 and float(r["alpha"]) in {0.0, 0.25, 0.5} and r["method"] in {"FedEraser-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"}:
            rows.append(
                [
                    r["method"].replace("-proxy", ""),
                    fnum(r["alpha"], 2),
                    fnum(r["mean_riskgap"], 3),
                    fnum(r["audit_pass_rate"], 2),
                    fnum(r["mean_model_residue"], 3),
                ]
            )
    return rows


def add_introduction(doc, raw):
    doc.add_heading("I. Introduction", level=1)
    para(doc, "Cross-border identity authentication has moved from a back-office compliance task to a live distributed-systems problem. A bank that onboards a mobile customer in one jurisdiction may need credit-bureau attributes, telecom-number evidence, sanctions screening, and historical fraud signals held by institutions in other jurisdictions. A multinational bank may also train a federated risk model across regional branches while each branch remains bound by local customer-due-diligence and data-transfer rules. Digital identity guidance stresses risk-based identity proofing and authentication [1], remote onboarding guidance requires accountable controls around non-face-to-face customer verification [2], and data-protection law gives data subjects forget request and transfer constraints that do not disappear when model training is federated [3]-[6].")
    para(doc, "This setting creates a practical failure mode that is rarely captured by standard federated unlearning evaluations. When a customer, regulator, or partner authority issues a forget request, the server must not only run an unlearning repair; it must establish that the request was authorized by the correct combination of authorities, that the request event did not reveal sensitive attributes through visible metadata, and that an auditor can later inspect evidence without receiving all underlying policy details. The issue is acute for cross-border credit and banking identity checks: the fact that a forget request concerns a foreign-worker flag, a personal-loan attribute, or a high-risk onboarding channel can itself be sensitive.")
    para(doc, "Existing schemes have three intrinsic limitations in this scenario. First, federated learning and early federated unlearning methods are optimizer-centric: they define how to update a model after a forget request, but they do not bind the unlearning repair to a multi-authority authorization predicate [7]-[10]. Second, recent verifiable and privacy-preserving federated unlearning work strengthens accountability, dynamic participation, attack analysis, certified unlearning repair, or efficient adapters, yet the visible request surface still tends to be controlled by the server and therefore can leak request type or policy scope [11]-[16]. Third, classical machine unlearning, certified data repair, and SISA-style retraining clarify retained-set targets, but they do not explain how independent financial, telecom, and government authorities can authorize the same cross-border forget request without collapsing policy privacy into a single administrative domain [17]-[22].")
    para(doc, "We propose MA-ABE-FU, a multi-authority attribute-based, evidence-bound control plane for federated unlearning. The core idea is intentionally narrow: MA-ABE is not used as a learning-side repair mechanism. It is used to authorize and hide the policy capsule that triggers a learning-side unlearning repair module. The repair module can be oracle retraining, SISA-style retraining, cached replay, differentially private recovery, certified repair, or a future optimizer. MA-ABE-FU contributes the verifiable envelope around that repair: a padded update-family channel, two zero-knowledge proof relations, and a UCAP evidence object that binds policy commitments, channel transcript, model commitments, and residual-risk measurements.")
    add_fig(doc, "Fig. 1.tif", "Fig. 1. MA-ABE-FU control plane. Multi-authority policy authorization, learning-side unlearning repair, and UCAP audit evidence are separated but transcript-bound.")
    para(doc, "This paper makes three contributions.")
    contribs = [
        "A cross-border unlearning protocol that turns a forget request from a server-local optimizer call into a policy-authenticated transaction. The protocol contains four executable algorithms, two formal proof relations, and one audit object; in our experiments it lowers mean attribute-leakage AUC from 0.671-0.907 for representative FU baselines to 0.503, close to random guessing.",
        "A checkable security formulation for policy-authenticated update hiding. The proof separates cryptographic indistinguishability from empirical model residue, so the theorem does not overclaim complete disappearance of model influence; the residual term remains an explicit audit parameter rather than being hidden inside a negligible bound.",
        "A reproducible experimental package for non-IID federated unlearning in public financial identity datasets. Across 126 federated runs and 16 proxy-ablation cells, MA-ABE-FU preserves retained AUC at 0.787 on average, runs at 0.514x oracle retraining time for the learning-side repair, and reports measured control-plane overhead under both primitive proxy and BN254 pairing backends.",
    ]
    for item in contribs:
        doc.add_paragraph(item, style="List Number")


def add_related_work(doc):
    doc.add_heading("II. Related Work", level=1)
    para(doc, "Federated unlearning mechanisms. FedEraser replays historical client updates, FedRecovery uses privacy-preserving recovery, VeriFi adds verifiable unlearning, dynamic-participation methods account for client churn, FedMUA studies malicious unlearning attacks, Starfish-style certified repair improves client-level evidence, FedWiper uses adapters to lower unlearning repair cost, and recent TDSC work shows that camouflaged poisoning can survive naive unlearning repair [9]-[16]. Their common mechanism is to approximate or certify a retained-set model. The mechanism is powerful but incomplete for cross-border identity: the server still sees enough event metadata to classify forget requests unless the request channel is separately hidden. MA-ABE-FU therefore treats these methods as learning-plane modules and adds a control plane that hides the policy capsule and binds the repair transcript to audit evidence.")
    para(doc, "Machine unlearning theory. SISA, certified data repair, formal right-to-be-forgotten models, and recent surveys define retained-set targets and residual-risk metrics [17]-[22]. Their bottom-layer limitation is that request authority is usually modeled as a single requester or data controller. Cross-border credit and banking identity workflows instead need several authorities to jointly determine whether the requested scope is valid. MA-ABE-FU preserves the retained-set evaluation logic from this literature but adds a multi-authority authorization predicate before the unlearning repair is accepted.")
    para(doc, "Attribute-based access control. Attribute-based encryption introduced fine-grained policy encryption, CP-ABE, multi-authority issuance, decentralized ABE, and expressive LSSS policy encoding [23]-[29]. The 2025 multi-committee ABE and sharding-chain work shows that decentralized access control can scale under distributed administration [30]. The mechanism gap is that access control answers who may decrypt a resource, while federated unlearning must also prove that an unlearning repair was executed and that the visible request metadata did not expose the policy. MA-ABE-FU uses MA-ABE only as the authorization wrapper; unlearning repair remains in the learning plane.")
    para(doc, "Proof and audit infrastructure. Pairing-based primitives, Fiat-Shamir transformations, zero-knowledge proofs, concise pairing arguments, Bulletproofs, PLONK-style proof systems, transparent proofs, and append-only ledgers provide the evidence tools needed for portable verification [31]-[39]. Used alone, however, they certify statements supplied by the system designer. MA-ABE-FU defines the missing statement: an authorized policy capsule, a padded forget request transcript, an unlearning repair commitment, and a residual-risk report must all verify together.")
    table_caption(doc, "I", "Mechanism gap in existing approaches")
    add_table(
        doc,
        ["Line of work", "Bottom mechanism", "Why it fails in cross-border multi-authorization", "MA-ABE-FU response"],
        [
            ["FU replay/recovery", "cached updates, recovery, adapters", "server-visible request type and policy scope remain outside the optimizer", "hide request capsule and bind repair to UCAP evidence"],
            ["Verifiable FU", "proof of an unlearning computation", "proof target does not encode multi-authority authorization privacy", "add pi_auth and padded update-family channel"],
            ["Certified repair/SISA", "shards or repair certificates", "single-domain request semantics do not express independent authorities", "compile forget scope as LSSS multi-authority policy"],
            ["MA-ABE/access control", "policy-based decryption", "access authorization does not measure model residue", "decouple authorization from learning-side repair"],
            ["Ledger/ZK audit", "append-only evidence and proof systems", "statements can omit malicious-server observability", "bind channel, policy, model commitments, and RiskGap"],
        ],
        [1.15, 1.45, 2.25, 1.65],
        size=7.4,
    )


def add_threat_model(doc):
    doc.add_heading("III. System and Threat Model", level=1)
    para(doc, "Entities. The system contains clients C_1,...,C_n, an authentication server AS, attribute authorities AA_1,...,AA_m, a requester Rq, and an auditor Aud. Each AA_i controls a disjoint or partially overlapping attribute domain U_i and issues GID-bound keys. A forget policy P is represented as an LSSS matrix with row map rho.")
    formula(doc, "P = (M, rho),    M in Z_p^{ell x k},    rho: {1,...,ell} -> union_i U_i.                         (1)")
    para(doc, "Planes. The control plane authorizes, hides, records, and audits the forget event. The learning plane receives an approved repair scope and runs a selected FU optimizer. No step in the protocol assumes that MA-ABE decrypts model parameters or performs statistical repair.")
    formula(doc, "F = (P, scope, tau, cfg),    S_F = {z in D : phi_P(z)=1 and epoch(z) <= tau}.                     (2)")
    para(doc, "Server behavior is split into two cases. During ordinary training the server is honest-but-curious: it executes FedAvg-style aggregation, observes timestamps, envelope lengths, client identifiers allowed by the deployment, update norms, and public configuration, and tries to infer membership or sensitive attributes. During the forget request phase it may be malicious: it can classify padded envelopes, infer policy attributes from metadata, delay unlearning repair, weaken repair rounds, omit clients from repair, reorder or suppress audit records, and forge evidence objects.")
    para(doc, "Attacker boundary. The adversary may adaptively request transcripts, corrupt a bounded set of clients, choose equal-length challenge envelopes, and query audit verification. It may not corrupt all authorities needed to satisfy the challenge policy, break MA-ABE collusion resistance, forge AS signatures, find commitment collisions, or make a false statement accepted by the proof system. Weak or incomplete repair is not hidden by the theorem; it is detected by pi_rep and the empirical RiskGap report.")
    table_caption(doc, "II", "Notation and formal objects")
    add_table(
        doc,
        ["Symbol", "Meaning", "Security role"],
        [
            ["AA_i, U_i, GID", "authority i, its attribute domain, and global identity", "multi-authority key boundary"],
            ["P=(M,rho)", "LSSS policy with matrix M in Z_p^{ell x k} and row map rho", "hidden authorization predicate"],
            ["F=(P,scope,tau,cfg)", "forget request descriptor: policy, scope, epoch, repair configuration", "request semantics"],
            ["S_F", "records satisfying phi_P within scope and epoch tau", "forget scope for unlearning repair"],
            ["Pad, e_F, ch", "padding mechanism, padded envelope, and channel transcript", "request-type hiding surface"],
            ["CT_F, tag", "MA-ABE capsule and encrypted request tag", "policy-bearing ciphertext"],
            ["c_P,c_s,c_cfg,h_F", "commitments to policy, scope, configuration, and request handle", "audit binding"],
            ["theta_t, theta'_t", "model commitments before and after unlearning repair", "repair binding"],
            ["R, s_l2, tau_R", "residual report, L2 normalization scale, and audit threshold", "RiskGap calibration"],
            ["R_auth, w_auth", "authorization proof relation and witness", "policy satisfaction without policy disclosure"],
            ["R_rep, w_rep, T_rep", "repair-consistency relation, witness, and repair trace", "unlearning repair consistency"],
            ["pi_auth, pi_rep", "zero-knowledge proofs for R_auth and R_rep", "proof handles"],
            ["O_F", "UCAP evidence object with commitments, proofs, signature, and chain pointer", "audit-chain unit"],
            ["alpha,beta", "RiskGap weights for model residue and utility loss", "residual-risk calibration"],
        ],
        [0.95, 2.45, 2.75],
        size=6.7,
    )


def add_protocol(doc):
    doc.add_heading("IV. MA-ABE-FU Protocol", level=1)
    para(doc, "The construction exposes four algorithms. The first two establish a policy-authenticated forget request. The third performs evidence-bound repair through a learning-side module. The fourth verifies the evidence object. All ciphertexts, commitments, and proofs are bound to a deployment identifier, an epoch, and a repair configuration to prevent cross-run replay.")
    alg_box(doc, "Algorithm 1: SystemSetup(1^lambda, {AA_i}, U, cfg)", "1: For each authority AA_i, sample (mpk_i, msk_i) over attribute domain U_i.\n2: Issue GID-bound keys SK_{GID,i,S_i}; bind each key to authority i and global identity GID.\n3: AS samples signing key sk_AS and initializes h_0 = H(domain || cfg).\n4: Publish pp = ({mpk_i}, Com, Pad, ZK parameters, audit grammar, tau_R).\n5: Output pp to requesters, clients, AS, and auditors; keep {msk_i} local to authorities.")
    alg_box(doc, "Algorithm 2: ForgetRequest(pp, GID, P, scope, tau, cfg)", "1: Compile P to LSSS matrix (M,rho) and compute tag = H(GID || scope || tau || cfg || nonce).\n2: Generate CT_F = MA-ABE.Enc({mpk_i}, tag; P).\n3: Build public commitments c_P=H(P), c_s=H(scope), c_cfg=H(cfg), and request handle h_F=H(CT_F || c_s || tau).\n4: Generate pi_auth for relation R_auth without exposing the satisfying attribute set.\n5: Emit e_F = Pad(type=forget, CT_F, c_P, c_s, c_cfg, tau, pi_auth) through the ordinary update queue.")
    alg_box(doc, "Algorithm 3: EvidenceBoundRepair(pp, e_F, theta_t, clients)", "1: Check that e_F belongs to the padded update-family grammar and that CT_F is well formed.\n2: Verify pi_auth; if it fails, reject before learning-plane repair.\n3: Derive S_F from the authorized scope and call Repair(theta_t, D \\ S_F, cfg).\n4: Commit theta'_t, channel transcript ch, repair trace T_rep, and residual report R.\n5: Generate pi_rep for relation R_rep and append O_F to the UCAP hash chain.\n6: Output theta'_t, R, and O_F=<c_P,c_s,H(ch),H(theta_t),H(theta'_t),H(R),pi_auth,pi_rep,sigma_AS,h_prev>.")
    alg_box(doc, "Algorithm 4: AuditVerify(pp, O_F, tau_R)", "1: Verify sigma_AS and the predecessor hash h_prev.\n2: Verify pi_auth on public input (c_P,c_s,c_cfg,CT_F,h_F,epoch).\n3: Verify pi_rep on public input (H(ch),H(theta_t),H(theta'_t),H(R),cfg).\n4: Recompute or sample-check residual report R on held-out audit data.\n5: Accept iff the proof checks pass and RiskGap(R) <= tau_R.")
    para(doc, "The two proof relations are the cryptographic boundary between hidden authorization and observable learning repair. Let w_auth contain a satisfying attribute-key tuple and encryption randomness. Let w_rep contain the repair trace, retained-client commitments, and opening randomness for model and report commitments.")
    formula(doc, "R_auth = { (c_P,c_s,c_cfg,CT_F,h_F,epoch; w_auth) : MA-ABE.Open(CT_F,w_auth)=tag and phi_P(S_GID)=1 }.       (3)")
    formula(doc, "R_rep = { (H(ch),H(theta_t),H(theta'_t),H(R),cfg; w_rep) : T_rep follows Repair(cfg,D\\S_F) and commits to theta'_t,R }.  (4)")
    formula(doc, "pi_auth <- ZK.Prove(R_auth; w_auth),       pi_rep <- ZK.Prove(R_rep; w_rep).                         (5)")
    para(doc, "The server therefore cannot use plaintext policy parsing as the repair trigger. It observes a padded envelope family and commitments; policy satisfaction is checked through pi_auth. Conversely, pi_rep does not prove that all influence has vanished. It proves that the declared repair transcript is consistent with the committed configuration, after which empirical residual risk is checked separately.")
    formula(doc, "RiskGap(R) = |AUC_MIA - 0.5| + alpha * ||theta'_t - theta_oracle||_2 / s_l2 + beta * max(0,U_oracle - U_ret).        (6)")


def add_security(doc):
    doc.add_heading("V. Security Analysis", level=1)
    para(doc, "Definition 1, policy-authenticated update hiding. The adversary A receives pp and may query key-issue, transcript, repair, and audit oracles, except that it may not obtain enough challenge attributes to satisfy P*. A chooses two equal-public-shape events e_0 and e_1: e_0 is an ordinary training update and e_1 is a valid forget event under P*. The challenger samples b, returns View_b, and A wins if it outputs b'=b.")
    formula(doc, "Adv_pauh(A) = | Pr[b'=b] - 1/2 |.                                                              (7)")
    para(doc, "The game covers the malicious-server actions stated in the threat model. Metadata classification is the View_b distinction problem. Attribute inference is blocked by the MA-ABE capsule and pi_auth. Audit-chain forgery is reduced to signature and commitment failure. Weak repair is not modeled away; it must either violate pi_rep or appear as residual RiskGap.")
    para(doc, "Theorem 1. Assume selective collusion-resistant MA-ABE, programmable random-oracle commitments, computationally indistinguishable padding for equal-shape envelopes, EUF-CMA secure AS signatures, and knowledge soundness for pi_auth and pi_rep. For every PPT adversary A against Definition 1 there exist PPT adversaries B_1,...,B_5 such that")
    formula(doc, "Adv_pauh(A) <= Adv_MA-ABE(B_1) + Adv_RO(B_2) + Adv_PAD(B_3) + Adv_SIG(B_4) + Adv_ZK(B_5) + eps_R + negl(lambda).       (8)")
    para(doc, "The term eps_R is the configured residual-risk tolerance. It is explicit because model influence is empirical for the selected learning module. This separation is important: the theorem proves policy-authenticated update hiding and audit binding for the control plane, not unconditional disappearance of model influence by encryption.")
    add_fig(doc, "Fig. 2.tif", "Fig. 2. Policy-authenticated update-hiding game and reduction purpose. G0-G5 respectively bound commitment replacement, MA-ABE capsule indistinguishability, padded-channel request-type hiding, UCAP-chain unforgeability, and zero-knowledge proof soundness.")
    para(doc, "Proof sketch. Game G0 is the real experiment. G1 replaces challenge commitments with lazy-sampled random-oracle values. G2 swaps the challenge MA-ABE capsule with an independent tag under the same public shape. G3 replaces visible envelope metadata with the padded distribution. G4 rejects any new UCAP object without a valid AS signature and predecessor hash. G5 rejects false pi_auth or pi_rep statements by knowledge soundness. In G5 the adversary sees only padded metadata, random commitments, valid signatures, and residual-risk evidence, so the remaining distinguishing advantage is bounded by eps_R plus a negligible term. Appendix A gives the full reduction.")


def add_experiments(doc, meta, bafs_status):
    doc.add_heading("VI. Experimental Methodology", level=1)
    para(doc, "Data and partitioning. Experiments use two public financial identity datasets: UCI German Credit and UCI Bank Marketing [40], [41]. German Credit is partitioned into 12 clients and Bank Marketing into 24 clients. Client partitions are Dirichlet non-IID with alpha=0.35, and each dataset is evaluated over three random seeds and three forget ratios: 0.25, 0.50, and 1.00 of the policy-eligible records. The BAFS suite is implemented as an optional loader because it is the preferred public fraud benchmark for large-scale banking identity experiments [42]. In this run, no authenticated BAFS CSV was present locally; the package records this status and does not report BAFS metrics.")
    table_caption(doc, "III", "Datasets and federated partition settings")
    add_table(
        doc,
        ["Dataset", "Rows", "Encoded features", "Clients", "Rounds", "Forget policy"],
        [
            ["German Credit", meta["German Credit"]["instances"], meta["German Credit"]["encoded_features"], meta["German Credit"]["clients"], meta["German Credit"]["rounds"], meta["German Credit"]["forget_policy"]],
            ["Bank Marketing", meta["Bank Marketing"]["instances"], meta["Bank Marketing"]["encoded_features"], meta["Bank Marketing"]["clients"], meta["Bank Marketing"]["rounds"], meta["Bank Marketing"]["forget_policy"]],
            ["BAFS", "not reported", "loader-ready", "32 if present", "8 if present", "foreign/high-risk internet-origin forget requests when columns exist"],
        ],
        [0.95, 0.62, 0.78, 0.58, 0.58, 3.00],
        size=7.0,
    )
    para(doc, "Baselines. The learning-plane comparison includes SISA-Retrain, FedEraser-proxy, FedRecovery-proxy, Starfish-proxy, MA-ABE-FU, and Oracle-Retrain. Oracle-Retrain retrains on retained federated data and is used only as a reference. Because official implementations do not share a common data pipeline and model class, the package implements faithful proxies and reports a proxy-ablation table rather than hiding this limitation.")
    table_caption(doc, "IV", "Baselines and fairness boundary")
    add_table(
        doc,
        ["Method", "Family", "Implemented mechanism", "Fairness boundary"],
        [
            ["SISA-Retrain", "sharded retraining", "affected-shard retained retraining", "same features, seeds, and client partition"],
            ["FedEraser-proxy", "cached replay", "history replay with affected-client recalibration", "official same-dataset metrics unavailable"],
            ["FedRecovery-proxy", "recovery", "retained-client recovery with calibrated perturbation", "proxy for recovery behavior"],
            ["Starfish-proxy", "certified repair", "short retained-set repair proxy", "no private circuit reproduction"],
            ["MA-ABE-FU", "control plane", "same unlearning repair family plus policy capsule, padding, proofs, UCAP", "authorization contribution is isolated"],
        ],
        [1.0, 1.15, 2.15, 2.1],
        size=7.1,
    )
    para(doc, "Attacks and metrics. Membership inference follows loss-threshold attacks in which forget scope training records are members and held-out policy scope records are non-members [43], [44]. Request-type leakage measures a malicious server classifier that distinguishes forget envelopes from ordinary updates. Attribute leakage measures whether visible metadata reveals the sensitive policy attribute. Retained AUC is the primary utility metric because the banking data are imbalanced.")
    para(doc, f"Code availability. The reproducibility repository is available at {CODE_URL}. The submission package also contains exact CSV outputs, the validation script, the IEEE figure-redrawing script, dependency notes, and the BAFS status file. The BAFS loader expects files under public_data/bafs/; when no authenticated CSV is present, the runner writes a status report and skips third-dataset numeric claims.")


def add_results(doc, raw, ablation, crypto, riskgrid):
    doc.add_heading("VII. Results", level=1)
    doc.add_heading("A. Unlearning Repair Quality", level=2)
    table_caption(doc, "V", "Primary federated unlearning results averaged over seeds and forget ratios")
    add_table(
        doc,
        ["Dataset", "Method", "Ret. AUC", "MIA gap", "L2/oracle", "Runtime/oracle", "Attr. leak AUC"],
        primary_rows(raw),
        [0.95, 1.15, 0.65, 0.65, 0.72, 0.86, 0.82],
        size=7.1,
    )
    add_fig(doc, "Fig. 3.tif", "Fig. 3. Unlearning repair quality under Dirichlet non-IID alpha=0.35. The figure aggregates forget ratios 0.25, 0.50, and 1.00; error bars are 95% confidence intervals over seeds; oracle retraining is the retained-set reference.")
    para(doc, "MA-ABE-FU matches the strongest repair proxies on retained AUC while adding the policy-authenticated envelope. Averaged over datasets, seeds, and forget ratios, MA-ABE-FU obtains retained AUC 0.787, compared with 0.786 for FedRecovery-proxy and 0.786 for Starfish-proxy. Its learning-side runtime is 0.514x oracle retraining on this convex model. FedEraser-proxy remains closest to oracle in parameter distance because cached replay is especially effective for logistic training; this is why the leakage and audit results are essential rather than auxiliary.")

    doc.add_heading("B. Malicious-Server Leakage", level=2)
    table_caption(doc, "VI", "Malicious-server observability and membership residue")
    add_table(
        doc,
        ["Method", "Request-type leak AUC", "Attribute leak AUC", "MIA gap"],
        leakage_rows(raw),
        [1.45, 1.35, 1.25, 0.95],
        size=7.4,
    )
    add_fig(doc, "Fig. 5.tif", "Fig. 5. Malicious-server leakage under Dirichlet non-IID alpha=0.35 and forget ratios 0.25, 0.50, and 1.00. Points and bars report mean AUC with 95% confidence intervals; 0.5 denotes random guessing. The figure can be read independently: MA-ABE-FU lowers attribute-leakage AUC to 0.503, while FU repair proxies average 0.826.")
    para(doc, "The malicious-server leakage figure is the central privacy result. FedEraser-proxy and FedRecovery-proxy expose both request type and policy attribute almost perfectly because their observable traces are shaped by the unlearning repair scope. Starfish-proxy weakens attribute inference to 0.671 but still leaves a visible repair signature. MA-ABE-FU moves attribute leakage to 0.503, an absolute gap of 0.398 against FedEraser-proxy, 0.405 against FedRecovery-proxy, and 0.168 against Starfish-proxy. Its request-type leakage remains 0.615 rather than 0.5, which indicates that traffic padding reduces but does not completely hide timing and queue-load signals.")

    doc.add_heading("C. Control-Plane Overhead and Audit Sensitivity", level=2)
    table_caption(doc, "VII", "Measured cryptographic overhead under primitive proxy and BN254 pairing backends")
    add_table(
        doc,
        ["Backend", "Authorities", "Rows", "Capsule/pairing ms", "HMAC ms", "Sign ms", "Total ms"],
        crypto_rows(crypto),
        [1.25, 0.72, 0.55, 1.15, 0.70, 0.70, 0.75],
        size=7.0,
    )
    add_fig(doc, "Fig. 4.tif", "Fig. 4. Measured cryptographic overhead of the MA-ABE-FU control plane. Panel a compares primitive proxy and BN254 pairing totals on a log scale across LSSS policy rows and authority counts; panel b isolates audit-chain micro-costs.")
    para(doc, "The primitive proxy ranges from 28.2 ms at 4 policy rows to 331.7 ms at 48 rows. The BN254 pairing backend ranges from 596.0 ms to 2708.3 ms as policy rows and authorities increase. Audit-chain components remain small: HMAC, RSA-PSS verification, and hash append are minor relative to capsule/pairing work. A production implementation should therefore optimize policy compilation, batching, and pairing operations before spending effort on the hash-chain component.")
    table_caption(doc, "VIII", "Proxy calibration and RiskGap sensitivity")
    rows = proxy_rows(ablation) + [["-- RiskGap beta=2 --", "", "", "", "", ""]] + risk_rows(riskgrid)
    add_table(
        doc,
        ["Dataset/Method", "Variant/alpha", "AUC dev.", "MIA delta/RiskGap", "Runtime/pass", "Max dev./residue"],
        rows,
        [1.45, 1.05, 0.85, 1.15, 0.90, 1.05],
        size=6.8,
    )
    para(doc, "The proxy ablation supports fair comparison. FedEraser-proxy deviates from oracle by at most 0.0003 AUC on German Credit and 0.000001 on Bank Marketing when affected-client recalibration is enabled; disabling recalibration increases German Credit deviation by roughly two orders of magnitude. FedRecovery-proxy and Starfish-proxy deviate by 0.006-0.009 AUC, which bounds how much interpretation should be attached to small utility differences among repair methods.")
    para(doc, "RiskGap is best used as an audit triage score rather than a legal constant. With beta=2, alpha=0 treats privacy residue and utility loss as the primary automatic gate and passes all methods whose MIA gap is close to oracle. Alpha=0.25 exposes nonzero model distance and flags recovery-style methods for manual review. We therefore recommend a two-tier deployment: alpha=0 for automatic channel/privacy acceptance, followed by alpha=0.25 residual review whenever the authority demands model-distance accountability.")


def add_discussion_conclusion(doc):
    doc.add_heading("VIII. Discussion", level=1)
    para(doc, "Horizontal comparison. Federated unlearning algorithms and cryptographic authorization schemes solve complementary parts of the cross-border identity problem. FedEraser, FedRecovery, VeriFi, Starfish-style certified repair, FedWiper, and SISA-style retraining concentrate on the learning plane: their strength is utility preservation or closeness to retained-set retraining [9]-[18]. MA-ABE and decentralized ABE concentrate on the authorization plane: their strength is distributed policy enforcement [23]-[30]. MA-ABE-FU occupies the missing interface. It does not claim a better optimizer; it makes an optimizer usable when a forget request must be authorized by several authorities and hidden from a server that can profile the request surface.")
    para(doc, "Engineering tradeoffs. Four variables dominate deployment. Increasing the number of authorities improves administrative separation but raises pairing latency, as Fig. 4 shows for two, four, and six authorities. Increasing the forget ratio can improve compliance coverage but may make unlearning repair less stable under non-IID clients. Larger LSSS policies improve expressiveness but dominate control-plane latency; batching and policy-row minimization should therefore be applied before changing the audit chain. Finally, stronger model-distance gates in RiskGap increase accountability but can reject otherwise privacy-favorable repairs. A practical configuration is a two-tier gate: alpha=0 for fast privacy-channel acceptance, followed by alpha=0.25 manual review when model-distance accountability is required.")
    para(doc, "Adversarial extensions. Data poisoning and camouflaged unlearning attacks can make an apparently valid unlearning repair amplify malicious influence [16]. MA-ABE-FU can be extended with client-update provenance, robust aggregation commitments, and poison-suspicion tags inside R so that pi_rep binds not only the repair trace but also the defense configuration. Authority collusion is the second natural extension. The current game excludes enough corrupted authorities to satisfy the challenge policy. A stronger deployment should use threshold authorities, authority-diversity constraints in rho, key-issuance transparency, and periodic cross-authority audits so that no single banking, telecom, or platform authority can silently authorize a high-risk forget request.")
    table_caption(doc, "IX", "Regulatory audit mapping")
    add_table(
        doc,
        ["Regime", "Relevant provisions", "MA-ABE-FU evidence", "Boundary"],
        [
            ["GDPR", "Arts. 5(2), 17, 30, 32; SCC transfer context", "O_F records policy, scope, channel, repair commitments, proofs, and RiskGap for accountability", "does not replace legal basis assessment"],
            ["China PIPL", "Arts. 38-40 cross-border transfer; Arts. 44 and 47 individual request rights", "multi-authority policy capsule separates authorization from server-visible metadata", "authority policy must reflect local compliance review"],
            ["CCPA/CPRA", "Cal. Civ. Code Sec. 1798.105 consumer request mechanism", "UCAP chain gives timestamped evidence for request handling and audit sampling", "statutory exemptions remain outside the protocol"],
        ],
        [0.85, 1.55, 2.55, 1.35],
        size=6.9,
    )
    para(doc, "Regulatory fit. The protocol is best understood as evidence support for compliance teams, not as a legal conclusion. GDPR accountability requires the controller to demonstrate how a statutory request was handled; PIPL emphasizes individual request rights and cross-border transfer governance; CCPA/CPRA requires operational response records for consumer requests [3]-[6]. MA-ABE-FU supplies a cryptographic audit trail for who authorized the forget request, which scope was committed, whether the request channel matched the padded grammar, which unlearning repair was executed, and what residual risk remained.")
    para(doc, "Limitations and future work. The FU baselines are faithful proxies, not official end-to-end reproductions; the ablation table reports the proxy error envelope to reduce bias. The BAFS loader is present, but authenticated data were not available in the local environment, so the manuscript does not claim BAFS numeric results. The current repair model is logistic and tabular; deep biometric or multimodal identity systems would require optimizer-specific unlearning repair and stronger membership-inference tests. Future work should add public BAFS runs once authenticated data are available, implement threshold-authority collusion resistance, integrate robust aggregation into pi_rep, and benchmark production pairing libraries with hardware acceleration.")
    doc.add_heading("IX. Conclusion", level=1)
    para(doc, "MA-ABE-FU introduces a policy-authenticated and evidence-bound control plane for federated unlearning in cross-border identity authentication. The protocol provides complete setup, forget request, unlearning repair, and audit algorithms; two formal zero-knowledge relations; a policy-authenticated update-hiding game; and a reduction proof that keeps empirical residual risk explicit. Experiments on non-IID public financial data show retained utility comparable to representative FU baselines while reducing malicious-server attribute leakage toward random guessing. The result is a practical and verifiable route for deploying federated unlearning where authorization, metadata privacy, and auditability are inseparable.")


def add_appendix(doc):
    doc.add_heading("Appendix A. Reduction Proof of Theorem 1", level=1)
    para(doc, "Let A be any PPT adversary that wins Definition 1 with non-negligible advantage. We define games G0 through G5 and construct simulators for each hop.")
    para(doc, "G0 to G1. G0 is the real experiment. In G1, every challenge commitment H(x) is replaced by a value sampled lazily from the random-oracle table. The simulator answers all repeated queries consistently. If A distinguishes the games, the simulator either finds two openings for one commitment or predicts an unqueried random-oracle value. This gives B_2, so the hop is bounded by Adv_RO(B_2)+negl(lambda).")
    para(doc, "G1 to G2. In G2, the challenge capsule CT_F is replaced with an encryption of an independent tag under the same public LSSS shape. Since A cannot obtain enough authority keys to satisfy the challenge policy and GID binding prevents cross-authority key pooling, any distinguisher yields B_1 against selective collusion-resistant MA-ABE. Hence the hop is bounded by Adv_MA-ABE(B_1).")
    para(doc, "G2 to G3. In G3, the visible channel transcript is replaced by the padded update-family distribution. The challenge pair has equal visible length and timing class by definition. A distinguisher for this hop is a distinguisher against the padding distribution, giving B_3 and bound Adv_PAD(B_3).")
    para(doc, "G3 to G4. In G4, any audit object that is not a previously returned object is rejected unless it carries a valid AS signature and predecessor-chain binding. A successful forgery gives B_4 against EUF-CMA signature security or a commitment collision already covered by the random-oracle hop. The additional probability is bounded by Adv_SIG(B_4).")
    para(doc, "G4 to G5. In G5, the challenger rejects if pi_auth or pi_rep accepts a false statement. A false accepted proof gives B_5 against knowledge soundness for the proof system. Thus the hop is bounded by Adv_ZK(B_5).")
    para(doc, "Final game. In G5 the adversary sees only padded metadata, random commitments, valid evidence objects, and the empirical residual report. The remaining distinction is exactly the audit tolerance eps_R induced by the selected learning module and RiskGap threshold. Summing the hops proves Theorem 1.")


def add_references(doc):
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFS, 1):
        para(doc, f"[{i}] {ref}", size=8.7)


def build_manuscript():
    raw = load_csv("federated_raw_v5.csv")
    ablation = load_csv("proxy_ablation_v5.csv")
    crypto = load_csv("crypto_overhead_v5.csv")
    riskgrid = load_csv("riskgap_sensitivity_v5.csv")
    meta = json.loads((REPRO / "validation_metadata_v5.json").read_text(encoding="utf-8"))
    bafs_status = json.loads((REPRO / "bafs_status_v5.json").read_text(encoding="utf-8"))

    doc = Document()
    setup_doc(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15.0)
    para(doc, AUTHORS, center=True, size=10.8)
    doc.add_heading("Abstract", level=1)
    para(doc, "Federated unlearning is usually evaluated as an unlearning repair problem: a server receives a forget request for a client or data subset and reports closeness to retained-set retraining. Cross-border identity authentication adds a second requirement. A forget event must be authorized by multiple independent authorities, hidden from a malicious server that can profile request metadata, and recorded as evidence that binds authorization, unlearning repair, and residual-risk measurements. This paper presents MA-ABE-FU, a policy-authenticated and evidence-bound control plane for federated unlearning. The protocol defines setup, forget request, evidence-bound unlearning repair, and audit verification algorithms; two zero-knowledge proof relations; and a formal policy-authenticated update-hiding game. The security proof reduces request observability and audit forgery to MA-ABE security, random-oracle commitments, padded-channel indistinguishability, signature unforgeability, and proof soundness, while keeping empirical model residue as an explicit tolerance. Experiments use non-IID federated partitions of German Credit and Bank Marketing, multiple forget ratios, membership inference, malicious-server request-type leakage, attribute leakage, FU proxy ablations, and measured primitive and BN254 pairing overhead. MA-ABE-FU preserves retained AUC comparable to representative FU repair baselines and lowers mean attribute-leakage AUC to 0.503, near random guessing. The result is not encryption-based disappearance of model influence; it is a verifiable control plane for authorized, metadata-aware federated forget requests.")
    para(doc, "Index Terms - Federated unlearning, multi-authority attribute-based encryption, cross-border identity authentication, zero-knowledge proof, membership inference, audit evidence.")
    add_introduction(doc, raw)
    add_related_work(doc)
    add_threat_model(doc)
    add_protocol(doc)
    add_security(doc)
    add_experiments(doc, meta, bafs_status)
    add_results(doc, raw, ablation, crypto, riskgrid)
    add_discussion_conclusion(doc)
    add_appendix(doc)
    add_references(doc)
    doc.save(OUT / "manuscript.docx")


def simple_doc(filename: str, lines):
    doc = Document()
    setup_doc(doc)
    for i, line in enumerate(lines):
        para(doc, line, bold=(i == 0), center=(i == 0), size=12 if i == 0 else 10.0)
    doc.save(OUT / filename)


def other_docs():
    simple_doc(
        "Title page.docx",
        [
            TITLE,
            "Jian Chen a,b, Sheng Peng b, Zhiming Cai b,c, Jiayin Qi d, Fu Mo e,*",
            "a Faculty of Finance and Economics, Guangdong University of Science and Technology, Dongguan, 523083, China",
            "b Faculty of Digital Science and Technology, Macau Millennium College, Macau, 999078, China",
            "c Institute for Data Engineering and Science, University of Saint Joseph, Macao, 999078, China",
            "d Cyberspace Institute of Advanced Technology, Guangzhou University, Guangzhou, 510000, China",
            "e Institute of Data Intelligence, Guangdong University of Science and Technology, Dongguan, 523083, China",
            f"*Corresponding author. Email: {EMAIL}. Tel: {TEL}.",
            "Funding",
            "This work was supported by the Macao FDCT (No. 002/2025/STT); the Chinese Academy of Engineering (No. 2025-XZ-08); and the 2025 Guangdong Provincial Project for Enhancing Research Capacity of Key Construction Disciplines (No. 2025ZDJS068).",
            "Declaration of competing interest",
            "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this work.",
            "Author contribution statement",
            "Jian Chen: Conceptualization, Methodology, Writing - Original Draft. Sheng Peng: Data Curation, Formal Analysis. Zhiming Cai: Validation, Formal Analysis. Jiayin Qi: Supervision, Project Administration, Writing - Review and Editing. Fu Mo: Supervision, Funding Acquisition, Resources, Writing - Review and Editing.",
        ],
    )
    simple_doc(
        "Cover Letter.docx",
        [
            "Cover Letter",
            "July 18, 2026",
            f"Editor-in-Chief, {TARGET}",
            "Dear Editor-in-Chief,",
            f"We submit \"{TITLE}\" as a regular research article.",
            "The manuscript presents a policy-authenticated and evidence-bound control plane for federated unlearning in cross-border identity authentication. It includes a complete protocol, formal security game, reduction proof, non-IID federated experiments, malicious-server observability tests, proxy-baseline ablations, and measured cryptographic overhead.",
            f"The code and reproducibility material are available at {CODE_URL}.",
            "The submission fits the scope of TDSC because it addresses dependable distributed systems, authentication, applied cryptographic authorization, auditability, and privacy-aware federated model maintenance.",
            "The manuscript is original, not under consideration elsewhere, and approved by all authors.",
            "Sincerely,",
            "Prof. Fu Mo",
            f"Email: {EMAIL}; Tel: {TEL}",
        ],
    )
    simple_doc(
        "Highlights.docx",
        [
            "Highlights",
            "MA-ABE-FU provides a policy-authenticated control plane for federated unlearning.",
            "The protocol defines four algorithms, two proof relations, and one formal update-hiding game.",
            "The reduction proof separates cryptographic guarantees from empirical model residue.",
            "Experiments cover non-IID clients, forget ratios, leakage attacks, and FU proxy ablations.",
            "The package reports measured primitive and BN254 pairing control-plane overhead.",
        ],
    )
    final_bio = next((p for p in (ROOT.parent / "final version").glob("Author*biography.docx")), None)
    if final_bio and final_bio.exists():
        shutil.copy2(final_bio, OUT / "Author biography.docx")
    else:
        simple_doc("Author biography.docx", ["Author Biography"] + BIOS)


def supplement():
    raw = load_csv("federated_raw_v5.csv")
    summary = load_csv("federated_summary_v5.csv")
    ablation = load_csv("proxy_ablation_v5.csv")
    crypto = load_csv("crypto_overhead_v5.csv")
    riskgrid = load_csv("riskgap_sensitivity_v5.csv")
    bafs_status = json.loads((REPRO / "bafs_status_v5.json").read_text(encoding="utf-8"))
    doc = Document()
    setup_doc(doc)
    para(doc, "Supplementary Material", bold=True, center=True, size=14)
    doc.add_heading("A. Artifact Manifest", level=1)
    para(doc, f"The reproducibility repository is {CODE_URL}. The submission archive contains run_validation_v5.py, redraw_ieee_figures_v5.py, prebuild_validation_v5.py, federated_raw_v5.csv, federated_summary_v5.csv, proxy_ablation_v5.csv, riskgap_sensitivity_v5.csv, crypto_overhead_v5.csv, validation_metadata_v5.json, term_audit_v5.csv, symbol_audit_v5.csv, figure_audit_v5.csv, reference_audit_v5.csv, and bafs_status_v5.json.")
    para(doc, f"BAFS status: available={bafs_status['available']}; expected files: {', '.join(bafs_status['expected_files'])}. No BAFS numeric row is used unless a listed CSV is present.")
    table_caption(doc, "S-I", "Federated result summary by ratio")
    add_table(
        doc,
        ["Dataset", "Ratio", "Method", "Ret. AUC", "MIA gap", "L2", "Type leak", "Attr leak"],
        [
            [r["dataset"], r["forget_ratio"], r["method"], fnum(r["retained_auc_mean"], 3), fnum(r["mia_gap_mean"], 3), fnum(r["l2_to_oracle_mean"], 3), fnum(r["type_leak_auc_mean"], 3), fnum(r["attribute_leak_auc_mean"], 3)]
            for r in summary
            if r["method"] in ["FedEraser-proxy", "FedRecovery-proxy", "Starfish-proxy", "MA-ABE-FU", "Oracle-Retrain"]
        ],
        [0.85, 0.45, 1.05, 0.62, 0.58, 0.55, 0.66, 0.66],
        size=6.5,
    )
    table_caption(doc, "S-II", "Proxy ablation cells")
    add_table(
        doc,
        ["Dataset", "Method", "Ablation", "AUC dev.", "MIA delta", "Runtime/oracle", "Max AUC dev."],
        [[r["dataset"], r["method"], r["ablation"], fnum(r["mean_abs_auc_delta_to_oracle"], 5), fnum(r["mean_mia_gap_delta_to_oracle"], 5), fnum(r["mean_runtime_to_oracle"], 3), fnum(r["max_abs_auc_delta_to_oracle"], 5)] for r in ablation],
        [0.75, 1.05, 1.35, 0.75, 0.75, 0.83, 0.83],
        size=6.4,
    )
    table_caption(doc, "S-III", "RiskGap grid excerpt")
    add_table(
        doc,
        ["alpha", "beta", "Method", "RiskGap", "Pass", "Residue", "Privacy gap", "Utility loss"],
        [[r["alpha"], r["beta"], r["method"], fnum(r["mean_riskgap"], 3), fnum(r["audit_pass_rate"], 2), fnum(r["mean_model_residue"], 3), fnum(r["mean_privacy_gap"], 3), fnum(r["mean_utility_loss"], 4)] for r in riskgrid if float(r["beta"]) in {0.0, 2.0} and float(r["alpha"]) in {0.0, 0.25, 1.0}],
        [0.5, 0.45, 1.05, 0.65, 0.45, 0.65, 0.68, 0.68],
        size=6.2,
    )
    table_caption(doc, "S-IV", "Cryptographic benchmark full table")
    add_table(
        doc,
        ["Backend", "Auth.", "Rows", "Capsule", "Pairing", "HMAC", "Sign", "Verify", "Hash", "Total"],
        [[r["backend"], r["authority_count"], r["policy_rows"], fnum(r["ma_abe_capsule_proxy_ms"], 2) if r["ma_abe_capsule_proxy_ms"] else "", fnum(r["bn254_pairing_ma_abe_ms"], 2) if r["bn254_pairing_ma_abe_ms"] else "", fnum(r["padded_envelope_hmac_ms"], 3), fnum(r["ucap_rsa_pss_sign_ms"], 3), fnum(r["ucap_rsa_pss_verify_ms"], 3), fnum(r["hash_chain_append_ms"], 3), fnum(r["total_control_plane_ms"], 2)] for r in crypto],
        [1.10, 0.42, 0.42, 0.63, 0.63, 0.52, 0.52, 0.52, 0.52, 0.63],
        size=5.9,
    )
    para(doc, f"Raw run count: {len(raw)} federated records. The CSV files in the reproducibility folder are the authoritative source for all plotted values.")
    doc.save(OUT / "Supplementary material.docx")


def bib_escape(text: str) -> str:
    return text.replace("\\", "\\textbackslash{}").replace("{", "\\{").replace("}", "\\}")


def write_references_bib():
    entries = []
    for i, ref in enumerate(REFS, 1):
        entries.append(f"@misc{{ref{i:02d},\n  note = {{{bib_escape(ref)}}}\n}}")
    (OUT / "references.bib").write_text("\n\n".join(entries) + "\n", encoding="utf-8")


def copy_reproducibility_files():
    if REPRO.exists():
        stale_names = {
            "federated_raw_v4.csv",
            "federated_summary_v4.csv",
            "crypto_overhead_v4.csv",
            "validation_metadata_v4.json",
            "run_validation_v4.py",
            "redraw_ieee_figures.py",
        }
        for item in REPRO.iterdir():
            if item.is_file() and (item.name in stale_names or "_v4" in item.name):
                item.unlink()
    shutil.copy2(ROOT / "run_validation_v5.py", REPRO / "run_validation_v5.py")
    shutil.copy2(ROOT / "redraw_ieee_figures_v5.py", REPRO / "redraw_ieee_figures_v5.py")
    shutil.copy2(ROOT / "prebuild_validation_v5.py", REPRO / "prebuild_validation_v5.py")
    README = f"""# MA-ABE-FU reproducibility package

Target manuscript: {TITLE}

Public repository: {CODE_URL}

## Contents

- `run_validation_v5.py`: federated partitioning, learning-plane baselines, leakage attacks, proxy ablations, RiskGap sensitivity, and cryptographic benchmarks.
- `redraw_ieee_figures_v5.py`: 600 dpi IEEE-style figure generation from CSV results.
- `prebuild_validation_v5.py`: terminology, symbol, figure, reference, and trace audit before manuscript construction.
- `build_submission_v5.py`: manuscript, supplementary material, title page, highlights, cover letter, and package builder.
- `submission_tdsc_v5/reproducibility/*.csv|*.json`: exact experiment outputs and validation records used in the manuscript.
- `submission_tdsc_v5/figure/Fig. *.pdf|*.tif`: vector PDFs and 600 dpi TIFF figures.
- `references.bib`: IEEE-formatted reference list exported for submission support.

## Reproduction

Install Python dependencies:

```bash
python -m pip install numpy pandas pillow cryptography py_ecc python-docx
```

Install Poppler, or set `PDFTOCAIRO` to the absolute path of `pdftocairo`, before redrawing figures.

Place public datasets under `public_data/`:

- `public_data/german.data`
- `public_data/bank/bank-full.csv`
- optional BAFS CSV files under `public_data/bafs/`

Run:

```bash
python run_validation_v5.py
python redraw_ieee_figures_v5.py
python prebuild_validation_v5.py
python build_submission_v5.py
```

The scripts regenerate all manuscript tables, exactly five main figures as vector PDFs plus 600 dpi TIFF files, the validation records, and the manuscript package. If BAFS files are absent, the runner writes `bafs_status_v5.json` and skips BAFS numeric results.
"""
    (OUT / "README_submission.md").write_text(README, encoding="utf-8")


def build_local_repro_repo():
    if REPRO_REPO.exists():
        shutil.rmtree(REPRO_REPO)
    (REPRO_REPO / "results").mkdir(parents=True)
    (REPRO_REPO / "figures").mkdir(parents=True)
    (REPRO_REPO / "submission_tdsc_v5" / "reproducibility").mkdir(parents=True)
    (REPRO_REPO / "submission_tdsc_v5" / "figure").mkdir(parents=True)
    shutil.copy2(ROOT / "run_validation_v5.py", REPRO_REPO / "run_validation_v5.py")
    shutil.copy2(ROOT / "redraw_ieee_figures_v5.py", REPRO_REPO / "redraw_ieee_figures_v5.py")
    shutil.copy2(ROOT / "prebuild_validation_v5.py", REPRO_REPO / "prebuild_validation_v5.py")
    shutil.copy2(ROOT / "build_submission_v5.py", REPRO_REPO / "build_submission_v5.py")
    shutil.copy2(OUT / "references.bib", REPRO_REPO / "references.bib")
    for name in [
        "federated_raw_v5.csv",
        "federated_summary_v5.csv",
        "proxy_ablation_v5.csv",
        "riskgap_sensitivity_v5.csv",
        "crypto_overhead_v5.csv",
        "validation_metadata_v5.json",
        "bafs_status_v5.json",
        "figure_manifest_v5.csv",
        "term_audit_v5.csv",
        "symbol_audit_v5.csv",
        "figure_audit_v5.csv",
        "reference_audit_v5.csv",
        "trace_audit_v5.csv",
        "prebuild_validation_v5.json",
    ]:
        src = REPRO / name
        if src.exists():
            shutil.copy2(src, REPRO_REPO / "results" / name)
            shutil.copy2(src, REPRO_REPO / "submission_tdsc_v5" / "reproducibility" / name)
    for fig in sorted(FIG.glob("Fig. *.*")):
        shutil.copy2(fig, REPRO_REPO / "figures" / fig.name)
        shutil.copy2(fig, REPRO_REPO / "submission_tdsc_v5" / "figure" / fig.name)
    (REPRO_REPO / "requirements.txt").write_text("numpy\npandas\npillow\ncryptography\npy_ecc\npython-docx\n", encoding="utf-8")
    (REPRO_REPO / "README.md").write_text((OUT / "README_submission.md").read_text(encoding="utf-8"), encoding="utf-8")
    (REPRO_REPO / "DATASETS.md").write_text(
        """# Dataset placement

The experiments use UCI German Credit and UCI Bank Marketing. The submission archive contains the local copies used for the reported run. For a public repository, place the same files as follows:

- `public_data/german.data`
- `public_data/bank/bank-full.csv`

The optional Bank Account Fraud Dataset Suite is evaluated only when one of these files is present:

- `public_data/bafs/Base.csv`
- `public_data/bafs/Variant I.csv`
- `public_data/bafs/Variant II.csv`
- `public_data/bafs/Variant III.csv`
- `public_data/bafs/Variant IV.csv`
- `public_data/bafs/Variant V.csv`

If no BAFS CSV is present, `run_validation_v5.py` writes `bafs_status_v5.json` and does not generate third-dataset metrics.
""",
        encoding="utf-8",
    )
    (REPRO_REPO / "PUBLISH.md").write_text(
        """# Repository publication

The directory is ready to publish as the reproducibility repository for the manuscript.

```bash
git init -b main
git add .
git commit -m "Add MA-ABE-FU reproducibility package"
git remote add origin https://github.com/Kent919/MA-ABE-FU.git
git push -u origin main
```

If the remote already exists, replace the `git remote add` command with:

```bash
git remote set-url origin https://github.com/Kent919/MA-ABE-FU.git
```

The repository contains scripts, exact CSV/JSON outputs, five vector PDF figures, five 600 dpi TIFF figures, and the reference export. Public datasets should remain outside version control unless redistribution rights are clear.
""",
        encoding="utf-8",
    )


def package_zip():
    zip_path = ROOT / "MA-ABE-FU_TDSC_submission_package_v5.zip"
    if zip_path.exists():
        zip_path.unlink()
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
        for f in OUT.rglob("*"):
            if f.is_file() and f.name != ".DS_Store":
                z.write(f, f.relative_to(OUT.parent))
        for f in REPRO_REPO.rglob("*"):
            if f.is_file() and f.name != ".DS_Store":
                z.write(f, f.relative_to(ROOT))
    return zip_path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    FIG.mkdir(parents=True, exist_ok=True)
    REPRO.mkdir(parents=True, exist_ok=True)
    required = [
        REPRO / "federated_raw_v5.csv",
        REPRO / "federated_summary_v5.csv",
        REPRO / "proxy_ablation_v5.csv",
        REPRO / "riskgap_sensitivity_v5.csv",
        REPRO / "crypto_overhead_v5.csv",
        REPRO / "validation_metadata_v5.json",
        REPRO / "bafs_status_v5.json",
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("Run run_validation_v5.py before building the submission package: " + ", ".join(missing))
    if len(list(FIG.glob("Fig. *.tif"))) != 5:
        raise RuntimeError("Run redraw_ieee_figures_v5.py before building; exactly five main TIFF figures are required.")
    copy_reproducibility_files()
    build_manuscript()
    other_docs()
    supplement()
    write_references_bib()
    build_local_repro_repo()
    zip_path = package_zip()
    print(zip_path)


if __name__ == "__main__":
    main()
